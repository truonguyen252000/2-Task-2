import streamlit as st
st.set_page_config(page_title="Power Quality Analysis", layout="wide", initial_sidebar_state="expanded")

st.markdown("""
<style>

/* Global styles */
.stApp {
    background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%);
    font-family: "Sitka Display Semibold", "Sitka Display", "Segoe UI", Roboto, sans-serif;
    text-align: center;
    font-size: 18px;
}

/* Sidebar styling */
section[data-testid="stSidebar"] {
    background: linear-gradient(180deg, #2c3e50 0%, #34495e 100%);
    padding: 2rem 1rem;
}
section[data-testid="stSidebar"] * {
    color: #ecf0f1 !important;
}
section[data-testid="stSidebar"] h1,
section[data-testid="stSidebar"] h2,
section[data-testid="stSidebar"] h3 {
    color: #3498db !important;
    font-weight: 700;
    margin-bottom: 1rem;
}

/* Main content area */
.main .block-container {
    padding: 2rem 3rem;
    max-width: 100%;
}

/* Title styling */
.main-title {
    font-size: 2.5rem;
    font-weight: 800;
    color: #2c3e50;
    text-align: center;
    margin-bottom: 2rem;
    text-shadow: 2px 2px 4px rgba(0,0,0,0.1);
}

/* Card styling */
.stat-card {
    background: white;
    padding: 1.5rem;
    border-radius: 10px;
    box-shadow: 0 4px 6px rgba(0,0,0,0.1);
    margin-bottom: 1rem;
}
.stat-number {
    font-size: 2.5rem;
    font-weight: 700;
    color: #3498db;
}
.stat-label {
    font-size: 1rem;
    color: #7f8c8d;
    margin-top: 0.5rem;
}

/* Button styling */
.stButton > button {
    width: 100%;
    background: linear-gradient(90deg, #3498db 0%, #2980b9 100%);
    color: white;
    border: none;
    border-radius: 8px;
    padding: 0.75rem 1.5rem;
    font-weight: 600;
    font-size: 1rem;
    box-shadow: 0px 4px 10px rgba(52, 152, 219, 0.3);
    transition: all 0.3s ease;
}
.stButton > button:hover {
    background: linear-gradient(90deg, #2980b9, #3498db);
    transform: translateY(-2px);
    box-shadow: 0px 6px 14px rgba(52, 152, 219, 0.4);
}

/* Input styling */
.stNumberInput input,
.stTextInput input {
    background-color: #34495e;
    border: 2px solid #4a5f7f;
    border-radius: 6px;
    color: white !important;
    padding: 0.5rem;
}
.stSelectbox select {
    background-color: #34495e;
    border: 2px solid #4a5f7f;
    border-radius: 6px;
    color: white;
}

/* Checkbox */
.stCheckbox {
    color: white !important;
}

/* Expander styling */
.streamlit-expanderHeader {
    background: linear-gradient(90deg, #ecf0f1 0%, #bdc3c7 100%);
    border-radius: 8px;
    font-weight: 600;
    color: #2c3e50 !important;
}

/* DataFrame */
.stDataFrame {
    border-radius: 10px;
    overflow: hidden;
    box-shadow: 0 4px 6px rgba(0,0,0,0.1);
}

/* Info boxes */
.stAlert {
    border-radius: 8px;
    border-left: 4px solid #3498db;
}

/* Divider */
hr {
    margin: 2rem 0;
    border: none;
    border-top: 2px solid #bdc3c7;
}

/* ⬇️ Download section */
.download-section {
    background: white;
    padding: 1.5rem;
    border-radius: 10px;
    box-shadow: 0 4px 6px rgba(0,0,0,0.1);
    margin: 1rem 0;
}

/* 📁 File uploader fix */
/* Remove white background of file uploader dropzone */
section[data-testid="stSidebar"] [data-testid="stFileUploaderDropzone"] {
    background-color: transparent !important;  /* loại bỏ nền trắng */
    border: 2px dashed rgba(52, 152, 219, 0.6) !important; /* hoặc bỏ dòng này nếu không muốn viền */
    border-radius: 10px !important;
    color: #ecf0f1 !important;
    transition: all 0.3s ease;
}

/*  Hover effect để người dùng biết có thể kéo thả */
section[data-testid="stSidebar"] [data-testid="stFileUploaderDropzone"]:hover {
    border-color: #3498db !important;
    background-color: rgba(52, 152, 219, 0.1) !important; /* tùy chọn: nền mờ nhẹ khi hover */
}

/* Màu chữ mô tả trong vùng dropzone */
section[data-testid="stSidebar"] [data-testid="stFileUploaderInstructions"] {
    color: rgba(236, 240, 241, 0.8) !important;
}

/* Nút Browse */
section[data-testid="stSidebar"] [data-testid="stFileUploaderBrowseButton"] {
    background: linear-gradient(90deg, #3498db 0%, #2980b9 100%) !important;
    color: white !important;
    border: none !important;
    border-radius: 6px !important;
    font-weight: 600 !important;
    box-shadow: 0px 3px 6px rgba(52, 152, 219, 0.3);
}
section[data-testid="stSidebar"] [data-testid="stFileUploaderBrowseButton"]:hover {
    background: linear-gradient(90deg, #2980b9, #3498db) !important;
}

section[data-testid="stSidebar"] [data-testid="stFileUploader"] {
    color: #3498db !important;
}
section[data-testid="stSidebar"] [data-testid="stFileUploader"] label div {
    color: #2c3e50 !important;
    font-weight: 800;
}
section[data-testid="stSidebar"] [data-testid="stFileUploader"] small {
    color: #2980b9 !important;
}
section[data-testid="stSidebar"] [data-testid="stFileUploader"] button {
    background: linear-gradient(90deg, #3498db 0%, #2980b9 100%) !important;
    color: white !important;
    border: none !important;
    border-radius: 6px !important;
    font-weight: 600 !important;
    box-shadow: 0px 3px 6px rgba(52, 152, 219, 0.3);
}
[data-testid="stFileUploader"] small {
    color: #2c3e50 !important;
    font-weight: 500 !important;
}

/* ✨ Improve text clarity */
body, p, label, span, div {
    letter-spacing: 0.3px;
    line-height: 1.5;
}


/* Tab styling - make text larger and bold */
.stTabs [data-baseweb="tab-list"] button [data-testid="stMarkdownContainer"] p {
    font-size: 1.3rem !important;
    font-weight: 700 !important;
    color: #2c3e50 !important;
}

.stTabs [data-baseweb="tab-list"] button {
    padding: 1rem 1.5rem !important;
}

.stTabs [data-baseweb="tab-list"] button[aria-selected="true"] [data-testid="stMarkdownContainer"] p {
    color: #3498db !important;
}          
</style>
""", unsafe_allow_html=True)


import pandas as pd
import numpy as np
import os
import glob
import io
import tempfile
import matplotlib.pyplot as plt
import warnings
from openpyxl.styles import Alignment
from reportlab.lib.pagesizes import A4, portrait
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
import re
import datetime
warnings.filterwarnings("ignore")
warnings.filterwarnings("ignore")


def clean_data(data, time_col_name="Time [UTC]"):
    if time_col_name not in data.columns:
        return data, "No time column found"
    
    df = data.copy()
    df[time_col_name] = pd.to_datetime(df[time_col_name], errors="coerce")
    df = df.dropna(subset=[time_col_name]).sort_values(time_col_name).reset_index(drop=True)
    
    if df.empty:
        return df, "Empty after cleaning"
    
    start, end = df[time_col_name].iloc[0], df[time_col_name].iloc[-1]
    start_is_midnight = start.time() == pd.Timestamp("00:00").time()
    end_is_2350 = end.time().hour == 23 and end.time().minute == 50
    
    log_msg = []
    
    if start_is_midnight and end_is_2350:
        df_cut = df.copy()
        log_msg.append("✅ Data already starts at 00:00 and ends at 23:50")
    else:
        start_date = (start + pd.Timedelta(days=1)).normalize()
        end_date = (end.normalize() - pd.Timedelta(minutes=10))
        df_cut = df[(df[time_col_name] >= start_date) & (df[time_col_name] <= end_date)]
        if df_cut.empty:
            return df_cut, "Empty after cutting"
        log_msg.append(f"📅 Cut data: {start_date.date()} to {end_date.date()}, {len(df_cut)} rows")
    
    df_cut = df_cut.sort_values(time_col_name).reset_index(drop=True)
    
    df_cut["date"] = df_cut[time_col_name].dt.date
    
    all_dates = sorted(df_cut["date"].unique())
    days_to_remove = []
    days_report = []
    
    for date in all_dates:
        expected_times = pd.date_range(
            start=pd.Timestamp(date),
            end=pd.Timestamp(date) + pd.Timedelta(hours=23, minutes=50),
            freq="10min"
        )
        
        actual_times = set(df_cut[df_cut["date"] == date][time_col_name])
        
        missing_count = len(set(expected_times) - actual_times)
        
        days_report.append(f"{date}: {missing_count}/144 missing")
        
        if missing_count > 15:
            days_to_remove.append(date)
    
    log_msg.append("\n📊 Daily Missing Report:")
    for report in days_report:
        log_msg.append(f"   {report}")
    
    if days_to_remove:
        df_cut = df_cut[~df_cut["date"].isin(days_to_remove)].copy()
        log_msg.append(f"\n🗑️ Removed {len(days_to_remove)} days with >15 missing timestamps")
        
        if df_cut.empty:
            return df_cut, "\n".join(log_msg) + "\n⚠️ No data remaining after filtering"
    else:
        log_msg.append("\n✅ All days have ≤15 missing timestamps")
    
    # Tính số ngày còn lại
    remaining_dates = sorted(df_cut["date"].unique())
    log_msg.append(f"\n📈 {len(remaining_dates)} valid days remaining")
    
    df_cut = df_cut.drop(columns=["date"]).reset_index(drop=True)
    
    if len(remaining_dates) > 0:
        all_times = []
        for date in remaining_dates:
            day_times = pd.date_range(
                start=pd.Timestamp(date),
                end=pd.Timestamp(date) + pd.Timedelta(hours=23, minutes=50),
                freq="10min"
            )
            all_times.extend(day_times)
        
        all_times = pd.DatetimeIndex(all_times)
        missing_times = all_times.difference(df_cut[time_col_name])
        
        if len(missing_times) > 0:
            log_msg.append(f"⚠️ Filling {len(missing_times)} missing timestamps in valid days")
            missing_df = pd.DataFrame({time_col_name: missing_times})
            df_full = pd.concat([df_cut, missing_df], ignore_index=True)
            df_full = df_full.sort_values(time_col_name).reset_index(drop=True)
            df_full = df_full.ffill().bfill()
        else:
            log_msg.append("✅ No missing timestamps in valid days")
            df_full = df_cut.copy()
        
        df_full["date"] = df_full[time_col_name].dt.date
        samples_per_day = df_full.groupby("date").size()
        days_incomplete = samples_per_day[samples_per_day < 144]
        if len(days_incomplete) > 0:
            log_msg.append(f"⚠️ {len(days_incomplete)} days still have <144 samples")
        else:
            log_msg.append(f"✅ All {len(samples_per_day)} days have exactly 144 samples")
        
        df_full = df_full.drop(columns=["date"])
    else:
        df_full = df_cut.copy()
    
    return df_full, "\n".join(log_msg)

# ===================== HÀM XỬ LÝ WORD REPORT =====================

def parse_excel_date(val):
    if pd.isna(val) or not val:
        return None
    if isinstance(val, (pd.Timestamp, datetime.datetime)):
        return val
    if isinstance(val, (int, float)) and val > 10000:
        try:
            return datetime.datetime.fromordinal(datetime.datetime(1900, 1, 1).toordinal() + int(val) - 2)
        except:
            return None
    if isinstance(val, str):
        for fmt in ("%d/%m/%Y", "%Y-%m-%d", "%d-%m-%Y", "%m/%d/%Y"):
            try:
                return datetime.datetime.strptime(val.strip(), fmt)
            except:
                continue
    return None

def format_cell_value(value, is_date=False, is_integer=False):
    if pd.isna(value):
        return ""
    if is_date and isinstance(value, datetime.datetime):
        return value.strftime("%d/%m/%Y")
    if is_integer:
        return str(int(value))
    if isinstance(value, (int, float)):
        return str(int(value)) if float(value).is_integer() else str(round(value, 2))
    return str(value)

def set_cell_format(cell, text, font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, color=None):
    cell.text = text
    for p in cell.paragraphs:
        p.alignment = alignment
        for run in p.runs:
            run.font.size = Pt(font_size)
            if color:
                run.font.color.rgb = color

def find_table_by_keyword(doc, keyword, search_rows=5):
    for tbl_idx, tbl in enumerate(doc.tables):
        for row_idx in range(min(search_rows, len(tbl.rows))):
            for cell in tbl.rows[row_idx].cells:
                if keyword.lower() in cell.text.strip().lower():
                    return tbl_idx, tbl, row_idx
    return None, None, None

def adjust_table_rows(table, needed_rows, header_rows=1):
    for r in range(header_rows, len(table.rows)):
        for c in table.rows[r].cells:
            c.text = ""
    
    current_rows = len(table.rows) - header_rows
    while current_rows < needed_rows:
        table.add_row()
        current_rows += 1
    while current_rows > needed_rows:
        table._tbl.remove(table.rows[-1]._tr)
        current_rows -= 1

def create_column_mapping(df_columns, word_headers):
    mapping = {}
    word_headers_lower = [h.strip().lower() for h in word_headers]
    for col in df_columns:
        name = str(col).lower()
        for idx, h in enumerate(word_headers_lower):
            if name == h or name in h or h in name or name.replace("%", "") == h.replace("%", ""):
                mapping[col] = idx
                break
    return mapping

def fill_harmonics_by_pdm_generic(doc, excel_path, sheet_name, keyword_indicator, table_name):
    table, bac_col_idx = None, None
    for tbl in doc.tables:
        found_bac_col = None
        has_indicator = False
        
        for row_idx, row in enumerate(tbl.rows):
            row_text = " ".join([c.text.strip() for c in row.cells]).lower()
            
            if keyword_indicator.lower() in row_text:
                has_indicator = True
                
            for cell_idx, cell in enumerate(row.cells):
                cell_lower = cell.text.strip().lower()
                if ("bậc" in cell_lower or "harmonic" in cell_lower) and len(cell.text.strip()) < 20:
                    found_bac_col = cell_idx
        
        if found_bac_col is not None and has_indicator:
            full_text = " ".join([" ".join([c.text for c in row.cells]) for row in tbl.rows]).lower()
            pdm_occurrences = full_text.count('%') + full_text.count('pđm') + full_text.count('pdm')
            
            if pdm_occurrences >= 5:
                table = tbl
                bac_col_idx = found_bac_col
                break

    if not table:
        return f"❌ Không tìm thấy bảng {table_name}"

    pdm_sections = []
    seen_pdm = set()
    
    for row_idx in range(min(len(table.rows), 50)):
        for cell_idx, cell in enumerate(table.rows[row_idx].cells):
            if cell_idx <= bac_col_idx:
                continue
            
            txt = cell.text.strip()
            m = re.search(r'(\d{1,3})\s*%', txt)
            
            if m:
                pdm_value = int(m.group(1))
                
                if pdm_value not in [10, 20, 30, 40, 50, 60, 70, 80, 90, 100] or pdm_value in seen_pdm:
                    continue
                
                first_max_col = None
                for check_row in range(row_idx, min(row_idx + 4, len(table.rows))):
                    for check_col in range(cell_idx, min(cell_idx + 20, len(table.rows[check_row].cells))):
                        cell_text = table.rows[check_row].cells[check_col].text.strip().lower()
                        if "max" in cell_text and len(cell_text) < 15:
                            first_max_col = check_col
                            break
                    if first_max_col is not None:
                        break
                
                if first_max_col is not None:
                    pdm_sections.append((pdm_value, first_max_col, row_idx))
                    seen_pdm.add(pdm_value)

    pdm_sections.sort(key=lambda x: x[0])

    df = pd.read_excel(excel_path, sheet_name=sheet_name, header=0)
    if isinstance(df.columns, pd.MultiIndex):
        df.columns = [" ".join([str(a).strip() for a in col if str(a).strip() and 'unnamed' not in str(a).lower()]) 
                      for col in df.columns]
    else:
        df.columns = [str(c).strip() for c in df.columns]

    pdm_col = next((c for c in df.columns 
                    if any(x in c.lower().replace(" ", "") 
                          for x in ["pdm", "%pđm", "pđm"])), 
                   df.columns[1] if len(df.columns) > 1 else df.columns[0])
    harm_col = next((c for c in df.columns 
                     if any(x in c.lower() for x in ["harm", "order", "bậc"])), None)
    if harm_col is None:
        for c in df.columns[:3]:
            try:
                if df[c].dropna().astype(str).str.strip().str.match(r'^\d+$').sum() >= 5:
                    harm_col = c
                    break
            except:
                continue

    df[pdm_col] = df[pdm_col].astype(str).str.replace(" ", "").str.lower()
    df_by_pdm = {}
    
    for pdm_value, _, _ in pdm_sections:
        patterns = [f"{pdm_value}%", f"{pdm_value}.0%", f"{pdm_value}%pdm", 
                   f"{pdm_value}%pđm", str(pdm_value), f"{pdm_value}.0"]
        mask = df[pdm_col].isin(patterns)
        
        try:
            mask = mask | df[pdm_col].str.contains(fr'^{pdm_value}(?:\.|%|\s|$)', regex=True, na=False)
        except:
            pass
        
        df_filtered = df[mask].copy().reset_index(drop=True)
        if len(df_filtered) > 0:
            df_by_pdm[pdm_value] = df_filtered

    try:
        harm_col_idx = df.columns.get_loc(harm_col) if harm_col else df.columns.get_loc(pdm_col)
        data_cols = list(df.columns[harm_col_idx + 1:harm_col_idx + 10])
    except:
        data_cols = list(df.columns[-9:]) if len(df.columns) >= 9 else list(df.columns)

    header_rows = sorted(set([r for _, _, r in pdm_sections]))
    data_regions = []
    
    for i, header_row in enumerate(header_rows):
        data_start = header_row + 2
        data_end = header_rows[i + 1] if i + 1 < len(header_rows) else len(table.rows)
        
        for j in range(data_start, len(table.rows)):
            if keyword_indicator.lower() in table.rows[j].cells[bac_col_idx].text.strip().lower():
                data_end = j + 1
                break
        
        applicable_pdm = [p for p, _, r in pdm_sections if r == header_row]
        data_regions.append((data_start, data_end, applicable_pdm))

    filled_stats = {p: 0 for p, _, _ in pdm_sections}
    
    for data_start, data_end, applicable_pdm in data_regions:
        for row_idx in range(data_start, data_end):
            row = table.rows[row_idx]
            bac_text = row.cells[bac_col_idx].text.strip()
            
            is_indicator = keyword_indicator.lower() in bac_text.lower()
            bac = None if is_indicator else (int(bac_text) if bac_text.isdigit() else None)
            
            if bac is None and not is_indicator:
                continue

            for pdm_value in applicable_pdm:
                df_pdm = df_by_pdm.get(pdm_value)
                if df_pdm is None:
                    continue

                start_col_max = next((col for p, col, _ in pdm_sections if p == pdm_value), None)
                if start_col_max is None:
                    continue

                excel_row = None
                if harm_col:
                    try:
                        if is_indicator:
                            matches = df_pdm[df_pdm[harm_col].astype(str).str.lower().str.contains(keyword_indicator.lower(), na=False)]
                        else:
                            matches = df_pdm[df_pdm[harm_col].astype(str).str.extract(r'(\d+)', expand=False).astype(float) == float(bac)]
                        
                        if len(matches) > 0:
                            excel_row = matches.iloc[0]
                    except:
                        pass
                
                if excel_row is None:
                    fallback_idx = row_idx - data_start
                    if 0 <= fallback_idx < len(df_pdm):
                        excel_row = df_pdm.iloc[fallback_idx]

                if excel_row is None:
                    continue

                for col_offset, col_name in enumerate(data_cols[:9]):
                    word_col = start_col_max + col_offset
                    if word_col >= len(row.cells):
                        break
                    
                    try:
                        value = excel_row[col_name]
                        if pd.isna(value):
                            continue
                        
                        val_text = (str(int(value)) if isinstance(value, (int, float)) and float(value).is_integer() 
                                   else f"{value:.2f}" if isinstance(value, (int, float)) 
                                   else str(value))
                        
                        set_cell_format(row.cells[word_col], val_text, font_size=5)
                        filled_stats[pdm_value] += 1
                    except:
                        continue

    total = sum(filled_stats.values())
    return f"✅ {table_name}: Điền {total} giá trị"
    
def process_word_report(excel_path, word_template_path, output_word_path, thresholds_dict=None):
    try:
        doc = Document(word_template_path)
        logs = []
        
        # PHẦN 0: Điền threshold placeholders
        if thresholds_dict:
            placeholder_map = {
                '{thd_th}': str(thresholds_dict.get('thdu', '')),
                '{tdd_th}': str(thresholds_dict.get('tddi', '')),
                '{pst_th}': str(thresholds_dict.get('pst', '')),
                '{plt_th}': str(thresholds_dict.get('plt', '')),
                '{uneg_th}': str(thresholds_dict.get('u_neg', '')),
                '{uh_th}': str(thresholds_dict.get('vh', '')),
                '{ih_th}': str(thresholds_dict.get('ch', ''))
            }
            
            # Replace in paragraphs
            for paragraph in doc.paragraphs:
                for placeholder, value in placeholder_map.items():
                    if placeholder in paragraph.text:
                        for run in paragraph.runs:
                            if placeholder in run.text:
                                run.text = run.text.replace(placeholder, value)
            
            # Replace in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for placeholder, value in placeholder_map.items():
                                if placeholder in paragraph.text:
                                    for run in paragraph.runs:
                                        if placeholder in run.text:
                                            run.text = run.text.replace(placeholder, value)
            
            logs.append(f"✅ Thresholds filled: THD={thresholds_dict.get('thdu')}%, TDD={thresholds_dict.get('tddi')}%, Pst={thresholds_dict.get('pst')}, Plt={thresholds_dict.get('plt')}, Uneg={thresholds_dict.get('u_neg')}%, VH={thresholds_dict.get('vh')}%, IH={thresholds_dict.get('ch')}%")
        # PHẦN 0.5: Điền thông tin thống kê dự án
        try:
            df_stats = pd.read_excel(excel_path, sheet_name="Daily_Pdm_Distribution")
            pdm_mw = thresholds_dict.get('pdm_mw', '') if thresholds_dict else ''

            df_data = df_stats[~df_stats["Date"].astype(str).str.contains("TOTAL|TỔNG", case=False, na=False)].copy()
            
            total_days = len(df_data)
            valid_days_count = len(df_data[df_data["Status"].astype(str).str.contains("Valid|✓", case=False, na=False)])
            
            pdm_cols = [col for col in df_stats.columns if '%' in str(col) and col not in ['Date', 'Index', 'Status']]
            
            if len(pdm_cols) > 0:
                total_row = df_stats[df_stats["Date"].astype(str).str.contains("TOTAL|TỔNG", case=False, na=False)]
                if len(total_row) > 0:
                    cols_above_50 = [col for col in pdm_cols if any(str(p) in col for p in ['50%', '60%', '70%', '80%', '90%', '100%'])]
                    samples_above_50 = int(total_row[cols_above_50].sum(axis=1).iloc[0]) if cols_above_50 else 0
                else:
                    cols_above_50 = [col for col in pdm_cols if any(str(p) in col for p in ['50%', '60%', '70%', '80%', '90%', '100%'])]
                    samples_above_50 = int(df_data[cols_above_50].sum().sum()) if cols_above_50 else 0
            else:
                samples_above_50 = 0
            
            valid_samples = thresholds_dict.get('valid_samples', 0) if thresholds_dict else 0
            total_samples = thresholds_dict.get('total_samples', 0) if thresholds_dict else 0
            
            stats_placeholder_map = {
                '{Pdm}': str(pdm_mw),
                '{valid_samples}': f"{valid_samples:,}",
                '{total_samples}': f"{total_samples:,}",
                '{sample_above_50}': f"{samples_above_50:,}",
                '{samples_above_50}': f"{samples_above_50:,}",  # Alternative spelling
                '{valid_days_count}': str(valid_days_count),
                '{total_days}': str(total_days)
            }
            
            # Replace in paragraphs
            for paragraph in doc.paragraphs:
                for placeholder, value in stats_placeholder_map.items():
                    if placeholder in paragraph.text:
                        for run in paragraph.runs:
                            if placeholder in run.text:
                                run.text = run.text.replace(placeholder, value)
            
            # Replace in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for placeholder, value in stats_placeholder_map.items():
                                if placeholder in paragraph.text:
                                    for run in paragraph.runs:
                                        if placeholder in run.text:
                                            run.text = run.text.replace(placeholder, value)
            
            logs.append(f"✅ Statistics filled: Pdm={pdm_mw}MW, Valid days={valid_days_count}/{total_days}, Samples≥50%={samples_above_50:,}/{total_samples:,}")
        except Exception as e:
            logs.append(f"⚠️ Statistics placeholders: {str(e)}")
        # PHẦN 1: Daily Distribution
        try:
            df_daily = pd.read_excel(excel_path, sheet_name="Daily_Pdm_Distribution")
            df_daily = df_daily.rename(columns={"Index": "STT", "Date": "Ngày"})
            df_daily = df_daily[~df_daily["Ngày"].astype(str).str.contains("TOTAL|TỔNG|Mean", case=False, na=False)].reset_index(drop=True)
            df_daily["Ngày"] = df_daily["Ngày"].apply(parse_excel_date)
            df_daily["STT"] = range(1, len(df_daily) + 1)

            daily_table = next((tbl for tbl in doc.tables if len(tbl.rows) > 0 and 
                               all(any(h.lower() in c.text.lower() for c in tbl.rows[0].cells) for h in ["stt", "ngày"])), None)

            if daily_table:
                header_cells = [c.text.strip() for c in daily_table.rows[0].cells]
                mapping = create_column_mapping(df_daily.columns, header_cells)
                adjust_table_rows(daily_table, len(df_daily))
                
                for i, row in df_daily.iterrows():
                    target_row = daily_table.rows[1 + i]
                    for col, idx in mapping.items():
                        if idx < len(target_row.cells):
                            text_val = format_cell_value(row[col], is_date=(col == "Ngày"), is_integer=(col == "STT"))
                            set_cell_format(target_row.cells[idx], text_val)
                
                logs.append(f"✅ Daily Distribution: {len(df_daily)} dòng")
            else:
                logs.append("⚠️ Không tìm thấy bảng Daily Distribution")
        except Exception as e:
            logs.append(f"❌ Daily Distribution: {str(e)}")

        # PHẦN 2: Summary tables
        summary_configs = [
            {"sheet": "THDu_summary", "keyword": "tổng số mẫu đo thd trong thời gian đo", 
             "rows": 4, "cols": 3, "row_range": (1, 5), "col_range": (1, 4), "name": "THD"},
            {"sheet": "Pst_summary", "keyword": "tổng số mẫu đo pst trong thời gian đo",
             "rows": 4, "cols": 3, "row_range": (1, 5), "col_range": (1, 4), "name": "Pst"},
            {"sheet": "Plt_24h_summary", "keyword": "tổng số mẫu đo plt trong thời gian đo",
             "rows": 4, "cols": 3, "row_range": (1, 5), "col_range": (1, 4), "name": "Plt"},
            {"sheet": "u0Avg_summary", "keyword": "tổng số mẫu đo uneg trong thời gian đo",
             "rows": 4, "cols": 1, "row_range": (1, 5), "col_range": (1, 2), "name": "Uneg"},
            {"sheet": "TDDi_summary", "keyword": "tổng số mẫu đo tdd trong thời gian đo",
             "rows": 4, "cols": 3, "row_range": (1, 5), "col_range": (1, 4), "name": "TDD"}
        ]

        for config in summary_configs:
            try:
                df_data = pd.read_excel(excel_path, sheet_name=config["sheet"])
                numeric_block = df_data.iloc[config["row_range"][0]:config["row_range"][1], 
                                              config["col_range"][0]:config["col_range"][1]].to_numpy(dtype=float)
                
                for tbl in doc.tables:
                    for r_idx, row in enumerate(tbl.rows):
                        if config["keyword"] in row.cells[0].text.strip().lower():
                            for i in range(config["rows"]):
                                for j in range(config["cols"]):
                                    set_cell_format(tbl.rows[r_idx + i].cells[j + 1], format_cell_value(numeric_block[i, j]))
                            logs.append(f"✅ {config['name']}: {config['rows']}×{config['cols']}")
                            break
                    else:
                        continue
                    break
            except Exception as e:
                logs.append(f"⚠️ {config['name']}: {str(e)}")

        # PHẦN 3: Kết quả đánh giá
        result_configs = [
            {"row_keyword": "tổng biến dạng sóng hài điện áp",
            "sheet": "THDu_summary", "percentage_row": 4, "status_row": 5, "name": "THD"},
            {"row_keyword": "tổng biến dạng sóng hài dòng điện",
            "sheet": "TDDi_summary", "percentage_row": 4, "status_row": 5, "name": "TDD"},
            {"row_keyword": "mức nhấp nháy điện áp ngắn hạn",
            "sheet": "Pst_summary", "percentage_row": 4, "status_row": 5, "name": "Pst"},
            {"row_keyword": "mức nhấp nháy điện áp dài hạn",
            "sheet": "Plt_24h_summary", "percentage_row": 4, "status_row": 5, "name": "Plt"},
            {"row_keyword": "mất cân bằng pha điện áp",
            "sheet": "u0Avg_summary", "percentage_row": 4, "status_row": 5, "name": "Uneg"}
        ]

        # Tìm bảng có header: "Thông tư 39/2015/TT-BCT | Kết quả đo | Đạt | Không đạt"
        result_table = None
        for tbl in doc.tables:
            if len(tbl.rows) == 0:
                continue
            
            header_row = tbl.rows[0]
            if len(header_row.cells) < 4:
                continue
            
            header_texts = [cell.text.strip().lower() for cell in header_row.cells]
            
            # Kiểm tra có đủ các cột cần thiết
            has_thongtu = any("39/2015" in txt or "thông tư" in txt for txt in header_texts)
            has_ketqua = any("kết quả" in txt for txt in header_texts)
            has_dat = any(txt == "đạt" or "đạt" in txt for txt in header_texts)
            has_khongdat = any("không đạt" in txt or "không" in txt for txt in header_texts)
            
            if has_thongtu and has_ketqua and has_dat and has_khongdat:
                result_table = tbl
                break

        if result_table is None:
            logs.append("⚠️ Không tìm thấy bảng Thông tư 39/2015/TT-BCT")
        else:
            header_row = result_table.rows[0]
            col_ketqua = next((i for i, c in enumerate(header_row.cells) if "kết quả" in c.text.lower()), 2)
            col_dat = next((i for i, c in enumerate(header_row.cells) if c.text.strip().lower() == "đạt"), 3)
            col_khongdat = next((i for i, c in enumerate(header_row.cells) if "không đạt" in c.text.lower()), 4)
            
            for config in result_configs:
                try:
                    df_result = pd.read_excel(excel_path, sheet_name=config["sheet"])

                    row_data = df_result.iloc[config["percentage_row"], 1:4]  
                    percentage_val = row_data.min()  
                    
                    status_val = "PASS" if percentage_val >= 95 else "FAIL"
                    
                    found = False
                    for row_idx, row in enumerate(result_table.rows[1:], start=1):
                        if len(row.cells) < col_khongdat + 1:
                            continue
                        
                        first_col_text = row.cells[0].text.strip().lower()
                        if config["row_keyword"].lower() in first_col_text:
                            found = True
                            
                            percentage_text = f"{format_cell_value(percentage_val)}%" if not pd.isna(percentage_val) else ""
                            set_cell_format(row.cells[col_ketqua], percentage_text)
                            
                            if status_val == "PASS":
                                set_cell_format(row.cells[col_dat], "x")
                                set_cell_format(row.cells[col_khongdat], "")
                            else:
                                set_cell_format(row.cells[col_dat], "")
                                set_cell_format(row.cells[col_khongdat], "x")
                            
                            logs.append(f"✅ Đã điền {config['name']}: {percentage_text} (MIN) - {'Đạt' if status_val == 'PASS' else 'Không đạt'}")
                            break
                    
                    if not found:
                        logs.append(f"⚠️ Không tìm thấy dòng '{config['row_keyword']}' cho {config['name']}")    
                except Exception as e:
                    logs.append(f"⚠️ Lỗi khi xử lý {config['name']}: {str(e)}")

        # PHẦN 4: Giá trị cực đại
        for keyword, sheet, name in [("giá trị thd cực đại pha a", "THD_Max_Values", "THD"),
                                      ("giá trị tdd cực đại pha a", "TDD_Max_Values", "TDD")]:
            try:
                df_max = pd.read_excel(excel_path, sheet_name=sheet)
                _, max_table, start_row = find_table_by_keyword(doc, keyword)
                
                if max_table:
                    for i in range(min(len(df_max), len(max_table.rows) - start_row)):
                        if len(max_table.rows[start_row + i].cells) > 3:
                            set_cell_format(max_table.rows[start_row + i].cells[1], format_cell_value(df_max.iloc[i, 1]))
                            set_cell_format(max_table.rows[start_row + i].cells[3], format_cell_value(df_max.iloc[i, 3]))
                    logs.append(f"✅ {name}: {i+1} dòng")
            except Exception as e:
                logs.append(f"⚠️ {name}: {str(e)}")
# PHẦN 5: Bảng thống kê Pst, Plt, Uneg theo %Pdm
        print("\n📋 PHẦN 5: Bảng thống kê theo %Pdm")
        
        def fill_stats_with_keys(sheet_name, keyword, name):
            try:
                df = pd.read_excel(excel_path, sheet_name=sheet_name)
                data_dict = {str(df.iloc[i, 0]).strip().lower().replace(" ", "").replace("\n", ""): 
                             [df.iloc[i, j] for j in range(1, len(df.columns))] for i in range(len(df))}
                
                _, table, start_row = find_table_by_keyword(doc, keyword)
                if not table:
                    return f"⚠️ {name}: Không tìm thấy bảng"
                
                filled = 0
                for row_idx in range(start_row + 2, len(table.rows)):
                    key = table.rows[row_idx].cells[0].text.strip().lower().replace(" ", "").replace("\n", "")
                    if not key:
                        continue
                    
                    for excel_key, values in data_dict.items():
                        if key in excel_key or excel_key in key:
                            for col_idx, value in enumerate(values):
                                if col_idx + 1 < len(table.rows[row_idx].cells) and value is not None:
                                    set_cell_format(table.rows[row_idx].cells[col_idx + 1], format_cell_value(value))
                            filled += 1
                            break
                
                return f"✅ {name}: {filled} dòng"
            except Exception as e:
                return f"⚠️ {name}: {e}"
        
        def fill_plt_table():
            try:
                df = pd.read_excel(excel_path, sheet_name="Plt_Statistics_24h")
                data_row = next((df.iloc[i] for i in range(len(df)) if "overall" in str(df.iloc[i, 0]).lower()), 
                                df.iloc[-1] if len(df) > 0 else None)
                
                if data_row is None:
                    return "⚠️ Plt: Không có dữ liệu"
                
                _, table, start_row = find_table_by_keyword(doc, "mức nhấp nháy điện áp dài hạn plt")
                if not table:
                    return "⚠️ Plt: Không tìm thấy bảng"
                
                # Tìm header row có "max avg min"
                header_row = next((r for r in range(start_row, min(start_row + 5, len(table.rows)))
                                  if "max" in " ".join([c.text.lower() for c in table.rows[r].cells])), None)
                
                if header_row is not None and header_row + 1 < len(table.rows):
                    values = [data_row.iloc[j] for j in range(1, min(10, len(data_row)))]
                    word_row = table.rows[header_row + 1]
                    filled = sum(1 for idx, val in enumerate(values) 
                                if idx + 1 < len(word_row.cells) and not pd.isna(val) and 
                                not set_cell_format(word_row.cells[idx + 1], format_cell_value(val)))
                    return f"✅ Plt: {len(values)} giá trị (3 pha × Max/AVG/Min)"
                else:
                    return "⚠️ Plt: Không tìm thấy header"
            except Exception as e:
                return f"⚠️ Plt: {e}"
        
        logs.append(fill_stats_with_keys("Power_Stats_Pst", "mức nhấp nháy điện áp ngắn hạn pst", "Pst"))
        logs.append(fill_plt_table())
        logs.append(fill_stats_with_keys("Power_Stats_Uneg", "cân bằng pha điện áp uneg", "Uneg"))

        # PHẦN 6: Harmonics by order
        for keyword, sheet, name in [("tổng số mẫu sóng hài riêng lẻ điện áp", "Voltage Harmonics by Order", "Sóng hài điện áp"),
                                      ("tổng số mẫu sóng hài riêng lẻ dòng điện", "Current Harmonics by Order", "Sóng hài dòng điện")]:
            try:
                _, table, start_row = find_table_by_keyword(doc, keyword)
                if not table:
                    logs.append(f"⚠️ {name}: Không tìm thấy")
                    continue
                
                df = pd.read_excel(excel_path, sheet_name=sheet, header=None).iloc[2:].reset_index(drop=True)
                df.columns = range(len(df.columns))
                df = df.rename(columns={1: 'Harmonic_Order'})
                
                header_row = start_row + 1
                bac_col = next((i for i, c in enumerate(table.rows[header_row].cells) if "bậc" in c.text.lower()), None)
                
                if bac_col is None:
                    logs.append(f"⚠️ {name}: Không tìm thấy cột 'Bậc'")
                    continue
                
                filled = 0
                for row_idx in range(header_row + 1, len(table.rows)):
                    row = table.rows[row_idx]
                    bac_text = row.cells[bac_col].text.strip()
                    if not bac_text.isdigit():
                        continue
                    
                    matching = df[df['Harmonic_Order'] == int(bac_text)]
                    if len(matching) == 0:
                        continue
                    
                    excel_row = matching.iloc[0]
                    
                    for offset, col_idx in enumerate(range(2, 8)):
                        if bac_col + 1 + offset < len(row.cells):
                            set_cell_format(row.cells[bac_col + 1 + offset], format_cell_value(excel_row.iloc[col_idx]))
                    
                    for offset, col_idx in enumerate(range(8, 11)):
                        if bac_col + 7 + offset < len(row.cells):
                            status = str(excel_row.iloc[col_idx]).strip().upper()
                            text = "ĐẠT" if status == "PASS" else "K.ĐẠT" if status == "FAIL" else ""
                            set_cell_format(row.cells[bac_col + 7 + offset], text, 
                                          color=RGBColor(255, 0, 0) if text == "K.ĐẠT" else None)
                    filled += 1
                
                logs.append(f"✅ {name}: {filled} bậc")
            except Exception as e:
                logs.append(f"⚠️ {name}: {str(e)}")

        # PHẦN 7: Harmonics by %Pdm
        logs.append(fill_harmonics_by_pdm_generic(doc, excel_path, "Voltage_Harmonics_by_%Pdm", "THD", "điện áp"))
        logs.append(fill_harmonics_by_pdm_generic(doc, excel_path, "Current_Harmonics_by_%Pdm", "TDD", "dòng điện"))

        doc.save(output_word_path)
        logs.append(f"✅ Hoàn tất! File đã lưu: {output_word_path}")
        
        return True, "\n".join(logs)
    
    except Exception as e:
        return False, f"❌ Lỗi xử lý Word: {str(e)}"


DEFAULT_OUTPUT_FOLDER = os.path.join(os.path.expanduser("~"), "PowerQuality_Output")

os.makedirs(DEFAULT_OUTPUT_FOLDER, exist_ok=True)
cols_pst = ["Pst1(Avg) []", "Pst2(Avg) []", "Pst3(Avg) []"]
cols_thdu = ["THD U1(AvgOn) [%]", "THD U2(AvgOn) [%]", "THD U3(AvgOn) [%]"]
cols_tddi = ["TDD I1(AvgOn) [%]", "TDD I2(AvgOn) [%]", "TDD I3(AvgOn) [%]"]
cols_plt = ["Plt1(Avg) []", "Plt2(Avg) []", "Plt3(Avg) []"]
cols_uneg = ["u-(Avg) [%]"]
thresholds = {
    "Pst1(Avg) []": 0.8, "Pst2(Avg) []": 0.8, "Pst3(Avg) []": 0.8,
    "THD U1(AvgOn) [%]": 3, "THD U2(AvgOn) [%]": 3, "THD U3(AvgOn) [%]": 3,
    "TDD I1(AvgOn) [%]": 3, "TDD I2(AvgOn) [%]": 3, "TDD I3(AvgOn) [%]": 3,
    "Plt1(Avg) []": 0.6, "Plt2(Avg) []": 0.6, "Plt3(Avg) []": 0.6,
    "u-(Avg) [%]": 3,
}

def update_thresholds(pst_val, plt_val, thdu_val, tddi_val, u_neg_val):
    global thresholds
    thresholds = {
        "Pst1(Avg) []": pst_val, "Pst2(Avg) []": pst_val, "Pst3(Avg) []": pst_val,
        "THD U1(AvgOn) [%]": thdu_val, "THD U2(AvgOn) [%]": thdu_val, "THD U3(AvgOn) [%]": thdu_val,
        "TDD I1(AvgOn) [%]": tddi_val, "TDD I2(AvgOn) [%]": tddi_val, "TDD I3(AvgOn) [%]": tddi_val,
        "Plt1(Avg) []": plt_val, "Plt2(Avg) []": plt_val, "Plt3(Avg) []": plt_val,
        "u-(Avg) [%]": u_neg_val,
    }
try:
    pdfmetrics.registerFont(TTFont('Times-Roman', 'times.ttf'))
except Exception:
    pass

def make_table(data, cols):
    df = data[cols]
    total = df.count()
    passed = pd.Series({c: int(((df[c] <= thresholds[c]) & df[c].notna()).sum()) for c in cols})
    failed = total - passed
    perc = (passed / total.replace(0, np.nan) * 100).round(2).fillna(0)
    status_per_col = perc.apply(lambda x: "PASS" if x >= 95 else "FAIL")
    group_status = "PASS" if (status_per_col == "PASS").all() else "FAIL"
    df_result = pd.DataFrame([
        total.astype(int),
        passed.astype(int),
        failed.astype(int),
        perc,
        status_per_col
    ], index=[
        "No. of samples",
        "No. of pass samples",
        "No. of fail samples",
        "Percentage of pass samples [%]",
        "Status"
    ])
    df_result.loc["Group Status"] = [group_status] * len(df_result.columns)
    return df_result

def make_voltage_harmonics_detail(data, threshold=1.5):
    orders = range(2, 41)
    phases = [("Phase A", "U1"), ("Phase B", "U2"), ("Phase C", "U3")]
    rows = []
    for h in orders:
        row = [h]
        above = []; perc_below = []; status = []
        for pha_name, prefix in phases:
            col = f"{prefix} h{h}(Avg) [%]"
            if col in data.columns:
                values = pd.to_numeric(data[col], errors="coerce")
                total = int(values.count())
                above_val = int((values > threshold).sum())
                below_val = int((values <= threshold).sum())
                perc = round(below_val / total * 100, 2) if total else 0.0
                stat = "PASS" if perc >= 95 else "FAIL"
            else:
                above_val = ""; perc = ""; stat = ""
            above.append(above_val); perc_below.append(perc); status.append(stat)
        row.extend(above); row.extend(perc_below); row.extend(status)
        rows.append(row)

    columns = [
        ("","Harmonic Order"),
        ("Total number of voltage harmonic samples > 1.5%", "Phase A"),
        ("Total number of voltage harmonic samples > 1.5%", "Phase B"),
        ("Total number of voltage harmonic samples > 1.5%", "Phase C"),
        ("Percentage of voltage harmonic samples ≤ 1.5% over total measured samples", "Phase A"),
        ("Percentage of voltage harmonic samples ≤ 1.5% over total measured samples", "Phase B"),
        ("Percentage of voltage harmonic samples ≤ 1.5% over total measured samples", "Phase C"),
        ("Assessment", "Phase A"),
        ("Assessment", "Phase B"),
        ("Assessment", "Phase C"),
    ]
    df_detail = pd.DataFrame(rows, columns=pd.MultiIndex.from_tuples(columns))
    return df_detail

def make_current_harmonics_detail(data, threshold=2):
    orders = range(2, 41)
    phases = [("Phase A", "I1"), ("Phase B", "I2"), ("Phase C", "I3")]
    rows = []
    for h in orders:
        row = [h]
        above = []; perc_below = []; status = []
        for pha_name, prefix in phases:
            col = f"{prefix} h{h}(Avg) [%]"
            if col in data.columns:
                values = pd.to_numeric(data[col], errors="coerce")
                total = int(values.count())
                above_val = int((values > threshold).sum())
                below_val = int((values <= threshold).sum())
                perc = round(below_val / total * 100, 2) if total else 0.0
                stat = "PASS" if perc >= 95 else "FAIL"
            else:
                above_val = ""; perc = ""; stat = ""
            above.append(above_val); perc_below.append(perc); status.append(stat)
        row.extend(above); row.extend(perc_below); row.extend(status)
        rows.append(row)
    columns = [
        ("","Harmonic Order"),
        ("Total number of current harmonic samples >2%", "Phase A"),
        ("Total number of current harmonic samples >2%", "Phase B"),
        ("Total number of current harmonic samples >2%", "Phase C"),
        ("Percentage of current harmonic samples ≤ 2% over total measured samples", "Phase A"),
        ("Percentage of current harmonic samples ≤ 2% over total measured samples", "Phase B"),
        ("Percentage of current harmonic samples ≤ 2% over total measured samples", "Phase C"),
        ("Assessment", "Phase A"),
        ("Assessment", "Phase B"),
        ("Assessment", "Phase C"),
    ]
    df_detail_cr = pd.DataFrame(rows, columns=pd.MultiIndex.from_tuples(columns))
    return df_detail_cr

def make_voltage_harmonics_max_table(data):
    orders = range(2, 41)
    phases = [("Phase A", "U1"), ("Phase B", "U2"), ("Phase C", "U3")]
    rows = []
    
    for h in orders:
        row = [h]
        for pha_name, prefix in phases:
            col = f"{prefix} h{h}(Avg) [%]"
            if col in data.columns:
                max_val = data[col].max()
                row.append(round(max_val, 3) if not np.isnan(max_val) else "")
            else:
                row.append("")
        rows.append(row)
    
    df = pd.DataFrame(rows, columns=["Voltage Harmonics by Order", "Phase A", "Phase B", "Phase C"])
    return df

def make_current_harmonics_max_table(data):
    orders = range(2, 41)
    phases = [("Phase A", "I1"), ("Phase B", "I2"), ("Phase C", "I3")]
    rows = []
    
    for h in orders:
        row = [h]
        for pha_name, prefix in phases:
            col = f"{prefix} h{h}(Avg) [%]"
            if col in data.columns:
                max_val = data[col].max()
                row.append(round(max_val, 3) if not np.isnan(max_val) else "")
            else:
                row.append("")
        rows.append(row)   
    df = pd.DataFrame(rows, columns=["Current Harmonics by Order", "Phase A", "Phase B", "Phase C"])
    return df

def make_thd_max_table(data):
    thd_cols = ["THD U1(AvgOn) [%]", "THD U2(AvgOn) [%]", "THD U3(AvgOn) [%]"]
    has_time = "Time [UTC]" in data.columns
    
    rows = []
    for i, thd_col in enumerate(thd_cols, 1):
        phase = chr(64 + i)  # A, B, C
        
        if thd_col in data.columns:
            max_thd = data[thd_col].max()
            max_idx = data[thd_col].idxmax()
            
            if "Ptot+(Avg) [W]" in data.columns:
                corresponding_power = data.loc[max_idx, "Ptot+(Avg) [W]"] / 1e6 
            else:
                corresponding_power = np.nan
            
            if has_time:
                timestamp = data.loc[max_idx, "Time [UTC]"]
                if pd.notna(timestamp):
                    timestamp_str = pd.to_datetime(timestamp).strftime("%Y-%m-%d %H:%M:%S")
                else:
                    timestamp_str = "N/A"
            else:
                timestamp_str = "N/A"
        else:
            max_thd = np.nan
            corresponding_power = np.nan
            timestamp_str = "N/A"
        
        rows.append([
            f"Maximum THD phase {phase} (% compared to rated voltage)",
            round(max_thd, 3) if not np.isnan(max_thd) else "N/A",
            f"Power plant output corresponding to maximum THD phase {phase} (MW)",
            round(corresponding_power, 2) if not np.isnan(corresponding_power) else "N/A",
            f"Timestamp of maximum THD phase {phase}",
            timestamp_str
        ]) 
    df_thd_max = pd.DataFrame(rows, columns=["Parameter", "Value", "Parameter 2", "Value 2", "Parameter 3", "Value 3"])
    return df_thd_max

def make_tdd_max_table(data):
    tdd_cols = ["TDD I1(AvgOn) [%]", "TDD I2(AvgOn) [%]", "TDD I3(AvgOn) [%]"]
    has_time = "Time [UTC]" in data.columns
    rows = []
    for i, tdd_col in enumerate(tdd_cols, 1):
        phase = chr(64 + i)  # A, B, C
        
        if tdd_col in data.columns:
            max_tdd = data[tdd_col].max()
            max_idx = data[tdd_col].idxmax()
            
            if "Ptot+(Avg) [W]" in data.columns:
                corresponding_power = data.loc[max_idx, "Ptot+(Avg) [W]"] / 1e6  # Convert to MW
            else:
                corresponding_power = np.nan
            
            if has_time:
                timestamp = data.loc[max_idx, "Time [UTC]"]
                if pd.notna(timestamp):
                    timestamp_str = pd.to_datetime(timestamp).strftime("%Y-%m-%d %H:%M:%S")
                else:
                    timestamp_str = "N/A"
            else:
                timestamp_str = "N/A"
        else:
            max_tdd = np.nan
            corresponding_power = np.nan
            timestamp_str = "N/A"
        
        rows.append([
            f"Maximum TDD phase {phase} (% compared to rated current)",
            round(max_tdd, 3) if not np.isnan(max_tdd) else "N/A",
            f"Power plant output corresponding to maximum TDD phase {phase} (MW)",
            round(corresponding_power, 2) if not np.isnan(corresponding_power) else "N/A",
            f"Timestamp of maximum TDD phase {phase}",
            timestamp_str
        ])  
    df_tdd_max = pd.DataFrame(rows, columns=["Parameter", "Value", "Parameter 2", "Value 2", "Parameter 3", "Value 3"])
    return df_tdd_max

def make_voltage_harmonics_by_pdm(data, bins, labels):
    harmonic_orders = range(2, 41)
    phases = [("Phase A", "U1"), ("Phase B", "U2"), ("Phase C", "U3")]
    thd_cols = ["THD U1(AvgOn) [%]", "THD U2(AvgOn) [%]", "THD U3(AvgOn) [%]"]
    rows_vh = []
    for (lower, upper), label in zip(bins, labels):
        if lower <= 0:
            mask = (data["Ptot+(Avg) [W]"] > 0) & (data["Ptot+(Avg) [W]"] <= upper)
        else:
            mask = (data["Ptot+(Avg) [W]"] > lower) & (data["Ptot+(Avg) [W]"] <= upper)
        df_bin = data.loc[mask]
        count = int(mask.sum())
        for h in harmonic_orders:
            row = [label, h]
            for pha_name, prefix in phases:
                col = f"{prefix} h{h}(Avg) [%]"
                if col in data.columns and count > 0:
                    mx = df_bin[col].max(); avg = df_bin[col].mean(); mn = df_bin[col].min()
                else:
                    mx = np.nan; avg = np.nan; mn = np.nan
                row.extend([mx, avg, mn])
            rows_vh.append(row)
        
        thd_row = [label, "THD"]
        for thd_col in thd_cols:
            if thd_col in data.columns and count > 0:
                mx = df_bin[thd_col].max(); avg = df_bin[thd_col].mean(); mn = df_bin[thd_col].min()
            else:
                mx = np.nan; avg = np.nan; mn = np.nan
            thd_row.extend([mx, avg, mn])
        rows_vh.append(thd_row)
        
    columns_vh = pd.MultiIndex.from_tuples([
        ("", "%Pdm"),
        ("", "Harmonic Order"),
        ("Phase A", "Max"), ("Phase A", "AVG"), ("Phase A", "Min"),
        ("Phase B", "Max"), ("Phase B", "AVG"), ("Phase B", "Min"),
        ("Phase C", "Max"), ("Phase C", "AVG"), ("Phase C", "Min"),
    ])
    df_vh = pd.DataFrame(rows_vh, columns=columns_vh)
    return df_vh.round(3)

def make_current_harmonics_by_pdm(data, bins, labels):
    harmonic_orders = range(2, 41)
    phases = [("Phase A", "I1"), ("Phase B", "I2"), ("Phase C", "I3")]
    tdd_cols = ["TDD I1(AvgOn) [%]", "TDD I2(AvgOn) [%]", "TDD I3(AvgOn) [%]"]
    rows_ch = []
    for (lower, upper), label in zip(bins, labels):
        if lower <= 0:
            mask = (data["Ptot+(Avg) [W]"] > 0) & (data["Ptot+(Avg) [W]"] <= upper)
        else:
            mask = (data["Ptot+(Avg) [W]"] > lower) & (data["Ptot+(Avg) [W]"] <= upper)
        df_bin = data.loc[mask]
        count = int(mask.sum())
        for h in harmonic_orders:
            row = [label, h]
            for pha_name, prefix in phases:
                col = f"{prefix} h{h}(Avg) [%]"
                if col in data.columns and count > 0:
                    mx = df_bin[col].max(); avg = df_bin[col].mean(); mn = df_bin[col].min()
                else:
                    mx = np.nan; avg = np.nan; mn = np.nan
                row.extend([mx, avg, mn])
            rows_ch.append(row)
        
        tdd_row = [label, "TDDi"]
        for tdd_col in tdd_cols:
            if tdd_col in data.columns and count > 0:
                mx = df_bin[tdd_col].max(); avg = df_bin[tdd_col].mean(); mn = df_bin[tdd_col].min()
            else:
                mx = np.nan; avg = np.nan; mn = np.nan
            tdd_row.extend([mx, avg, mn])
        rows_ch.append(tdd_row)
        
    columns_ch = pd.MultiIndex.from_tuples([
        ("", "%Pdm"),
        ("", "Harmonic Order"),
        ("Phase A", "Max"), ("Phase A", "AVG"), ("Phase A", "Min"),
        ("Phase B", "Max"), ("Phase B", "AVG"), ("Phase B", "Min"),
        ("Phase C", "Max"), ("Phase C", "AVG"), ("Phase C", "Min"),
    ])
    df_ch = pd.DataFrame(rows_ch, columns=columns_ch)
    return df_ch.round(3)

def make_daily_pdm_distribution(data, Pdm_max_val, data_raw=None):
    if "Time [UTC]" not in data.columns or "Ptot+(Avg) [W]" not in data.columns:
        return pd.DataFrame(), "", {}
    data_copy = data.copy()
    data_copy["Date"] = pd.to_datetime(data_copy["Time [UTC]"], errors='coerce').dt.date
    data_copy = data_copy.dropna(subset=["Date"])
    
    if len(data_copy) == 0:
        return pd.DataFrame(), "", {}
    total_samples_after_clean = len(data_copy)

    data_copy["%Pdm"] = (data_copy["Ptot+(Avg) [W]"] / Pdm_max_val * 100)
    pdm_ranges = [(i*10, (i+1)*10) for i in range(10)]
    pdm_labels = [f"{i*10}%" for i in range(1, 11)]
    dates = sorted(data_copy["Date"].unique())
    # Phân loại ngày đạt yêu cầu và không đạt
    valid_dates = []
    invalid_dates = []
    valid_dates_samples = {}
    valid_total_counts = [0] * len(pdm_ranges)  

    for date in dates:
        date_data = data_copy[data_copy["Date"] == date]
        if (date_data["%Pdm"] >= 50).any():
            valid_dates.append(date)
            valid_dates_samples[date] = len(date_data)
            
            # Tính total_counts chỉ cho valid days
            for i, (lower, upper) in enumerate(pdm_ranges):
                if i == 0:
                    count = len(date_data[(date_data["%Pdm"] > 0) & (date_data["%Pdm"] <= upper)])
                else:
                    count = len(date_data[(date_data["%Pdm"] > lower) & (date_data["%Pdm"] <= upper)])
                valid_total_counts[i] += count
        else:
            invalid_dates.append(date)

    rows = []
    total_counts = [0] * len(pdm_ranges) 

    for idx, date in enumerate(dates, 1):
        date_data = data_copy[data_copy["Date"] == date]
        row = [idx, date.strftime("%d/%m/%Y")]
        for i, (lower, upper) in enumerate(pdm_ranges):
            if i == 0:
                count = len(date_data[(date_data["%Pdm"] > 0) & (date_data["%Pdm"] <= upper)])
            else:
                count = len(date_data[(date_data["%Pdm"] > lower) & (date_data["%Pdm"] <= upper)])
            row.append(count)
            total_counts[i] += count
        
        if date in valid_dates:
            row.append(f"✓ Valid ({len(date_data)} samples)")
        else:
            row.append(f"✗ Invalid ({len(date_data)} samples)") 
        rows.append(row)

    # Thêm dòng Total
    total_row = ["", "TOTAL"] + total_counts + [""]
    rows.append(total_row)

    columns = ["Index", "Date"] + pdm_labels + ["Status"]
    df_distribution = pd.DataFrame(rows, columns=columns)

    if data_raw is not None and "Ptot+(Avg) [W]" in data_raw.columns:
        total_samples = len(data_raw)
    else:
        total_samples = len(data_copy)
    valid_samples = len(data_copy[data_copy["Ptot+(Avg) [W]"] > 0])

    # Dùng valid_total_counts
    samples_above_50 = sum(valid_total_counts[4:])  
    total_valid_days_samples = sum(valid_dates_samples.values())
    total_days = len(dates)
    valid_days_count = len(valid_dates)
    invalid_days_count = len(invalid_dates)

    summary_text = f"""**Recorded Power Statistics:**
    - Total measurement period: **{total_days} days**
    - **Valid days (≥ 50% Pdm): {valid_days_count} days** ✅
    - Invalid days (< 50% Pdm): {invalid_days_count} days ❌
    - **Total samples from valid days: {total_valid_days_samples:,} samples**
    - Total recorded samples: {valid_samples:,}/{total_samples:,} samples with power > 0 W (Ptot > 0 W)
    - Total samples with P ≥ 50% Pdm (from valid days only): **{samples_above_50:,} samples** 
    **Note:** Only data from valid days are used for further analysis."""

# Thông tin chi tiết về các ngày đạt yêu cầu
    detailed_info = {
        'total_days': total_days,
        'valid_days_count': valid_days_count,
        'invalid_days_count': invalid_days_count,
        'total_valid_days_samples': total_valid_days_samples,
        'valid_dates_samples': valid_dates_samples,
        'samples_above_50': samples_above_50, 
        'valid_samples': valid_samples,         
        'total_samples': total_samples         
    } 
    return df_distribution, summary_text, detailed_info

def plot_and_save(data, file_name, out_folder):
    groups = {"Pst": cols_pst, "THDu (%)": cols_thdu, "TDDi (%)": cols_tddi, "Plt": cols_plt}
    fig, axes = plt.subplots(4, 1, figsize=(12, 10), sharex=True)
    for ax, (title, cols) in zip(axes, groups.items()):
        for col in cols:
            if col in data.columns:
                ax.plot(pd.to_numeric(data[col], errors="coerce"), label=col)
        ax.set_ylabel(title, fontsize=10)
        ax.legend(fontsize=8)
        ax.grid(True)
    axes[-1].set_xlabel("Index (sample)", fontsize=10)
    plt.tight_layout()
    fig_path = os.path.join(out_folder, f"Output_{file_name}.png")
    plt.savefig(fig_path, dpi=150)
    plt.close()
    return fig_path

def save_table_to_excel(writer, df, group_name):
    df_copy = df.copy()
    df_copy.to_excel(writer, sheet_name=group_name + "_summary", index=True, startrow=1)
    worksheet = writer.sheets[group_name + "_summary"]
    last_col = len(df_copy.columns) + 1
    try:
        merge_range = f"A1:{chr(64+last_col)}1"
        worksheet.merge_cells(merge_range)
    except Exception:
        pass
    cell = worksheet.cell(row=1, column=1)
    cell.value = group_name
    cell.alignment = Alignment(horizontal='center', vertical='center')

def save_pdf(tables, out_pdf, base_name):
    PAGE_WIDTH, PAGE_HEIGHT = portrait(A4)
    LEFT_MARGIN = RIGHT_MARGIN = 36
    usable_width = PAGE_WIDTH - LEFT_MARGIN - RIGHT_MARGIN
    doc = SimpleDocTemplate(out_pdf, pagesize=portrait(A4),
                           leftMargin=LEFT_MARGIN, rightMargin=RIGHT_MARGIN)
    elements = []
    styles = getSampleStyleSheet()
    for style in styles.byName.values():
        style.fontName = 'Times-Roman'
    title = Paragraph(f"Summary Report - {base_name}", styles['Title'])
    elements.append(title)
    elements.append(Spacer(1, 12))

    for group_name, df in tables.items():
        elements.append(Paragraph(group_name, styles['Heading2']))
        elements.append(Spacer(1, 6))

        if isinstance(df.columns, pd.MultiIndex):
            nlevels = df.columns.nlevels
            header = []
            for level in range(nlevels):
                row = []
                for col in df.columns:
                    val = col[level] if isinstance(col, tuple) else ""
                    if level == 0 and val != "":
                        import textwrap
                        wrapped = "<br/>".join(textwrap.wrap(str(val), width=22))
                        val = Paragraph(f"<b><font size=11>{wrapped}</font></b>", ParagraphStyle('header', alignment=1, fontName='Times-Roman', leading=13))
                    elif level == 1 and val != "":
                        val = Paragraph(f"<b><font size=10>{val}</font></b>", ParagraphStyle('header', alignment=1, fontName='Times-Roman', leading=12))
                    else:
                        val = Paragraph(f"<font size=10>{val}</font>", ParagraphStyle('header', alignment=1, fontName='Times-Roman', leading=12))
                    row.append(val)
                header.append(row)
            data_table = header + df.astype(str).values.tolist()
            if df.index.name or df.index.names:
                for i, row in enumerate(data_table[nlevels:]):
                    data_table[nlevels + i] = [str(df.index[i])] + row
                for h in range(len(header)):
                    data_table[h] = [""] + data_table[h]
        else:
            data_table = [[df.index.name or "Metric"] + df.columns.tolist()]
            for idx, row in df.iterrows():
                data_table.append([str(idx)] + row.astype(str).tolist())

        n_cols = len(data_table[0])
        min_col_width = 50
        col_width = max(min_col_width, usable_width // n_cols)
        col_widths = [col_width] * n_cols
        total_width = sum(col_widths)
        if total_width > usable_width:
            scale = usable_width / total_width
            col_widths = [w * scale for w in col_widths]

        t = Table(data_table, hAlign="CENTER", colWidths=col_widths)
        style_list = [
            ('FONTNAME', (0, 0), (-1, -1), 'Times-Roman'),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('GRID', (0, 0), (-1, -1), 0.3, colors.black),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
        ]
        start_row = len(header) if isinstance(df.columns, pd.MultiIndex) else 1
        for r, row in enumerate(data_table[start_row:], start=start_row):
            for c, val in enumerate(row):
                if str(val).strip().upper() in ("PASS", "ĐẠT"):
                    style_list.append(('TEXTCOLOR', (c, r), (c, r), colors.green))
                elif str(val).strip().upper() in ("FAIL", "K.ĐẠT"):
                    style_list.append(('TEXTCOLOR', (c, r), (c, r), colors.red))
        t.setStyle(TableStyle(style_list))
        elements.append(t)
        elements.append(Spacer(1, 10))
    doc.build(elements)
    return out_pdf

# ======================= STREAMLIT UI =======================

# Initialize session state
if "run_processing" not in st.session_state:
    st.session_state.run_processing = False
if "processed_data" not in st.session_state:
    st.session_state.processed_data = {}
if "processing_complete" not in st.session_state:
    st.session_state.processing_complete = False

with st.sidebar:
    col1, col2, col3 = st.columns([0.5, 3, 0.5])
    with col2:
        logo_paths = [
            "logo.png",  
            os.path.join(os.path.dirname(__file__), "logo.png"),  
            os.path.join(os.getcwd(), "logo.png"), 
        ]
        
        logo_loaded = False
        for logo_path in logo_paths:
            if os.path.exists(logo_path):
                try:
                    st.image(logo_path, use_container_width=True)
                    logo_loaded = True
                    break
                except Exception:
                    continue

    st.markdown("#### Data Source")
    uploaded_files = st.file_uploader(
        "Upload Excel files (.xlsx)",
        type=["xlsx"],
        accept_multiple_files=True,
        help="Maximum 200MB per file"
    )
    
    folder_input = st.text_input(
        "Or folder path:",
        value="",
        help="Full path to folder with .xlsx files"
    )

    st.markdown("---")

    st.markdown("---")
    st.markdown("#### Output Settings")
    use_local_save = st.checkbox("Save to local folder", value=False,)
    if use_local_save:
        out_folder_input = st.text_input("Output folder:", value=DEFAULT_OUTPUT_FOLDER)
    else:
        out_folder_input = None
    
    
    st.markdown("---")
    st.markdown("#### Processing Options")
    enable_cleaning = st.checkbox("🧹 Clean data before analysis", value=True, 
                                   help="Cut first/last incomplete days, fill missing timestamps")
    
    st.markdown("---")
    st.markdown("#### Voltage Level & Thresholds")
    voltage_level = st.selectbox(
        "Select voltage level:",
        ["110kV", "22kV"],
        help="Different voltage levels have different threshold requirements"
    )
    
    if voltage_level == "110kV":
        threshold_pst = st.number_input("Pst threshold:", value=0.8, step=0.1, format="%.1f")
        threshold_plt = st.number_input("Plt threshold:", value=0.6, step=0.1, format="%.1f")
        threshold_thdu = st.number_input("THD U threshold (%):", value=3.0, step=0.5, format="%.1f")
        threshold_tddi = st.number_input("TDD I threshold (%):", value=3.0, step=0.5, format="%.1f")
        threshold_u_neg = st.number_input("u- threshold (%):", value=3.0, step=0.5, format="%.1f")
        threshold_vh = st.number_input("Voltage harmonic threshold (%):", value=1.5, step=0.1, format="%.1f")
        threshold_ch = st.number_input("Current harmonic threshold (%):", value=2.0, step=0.1, format="%.1f")
    else:  # 22kV
        threshold_pst = st.number_input("Pst threshold:", value=1.0, step=0.1, format="%.1f")
        threshold_plt = st.number_input("Plt threshold:", value=0.8, step=0.1, format="%.1f")
        threshold_thdu = st.number_input("THD U threshold (%):", value=5.0, step=0.5, format="%.1f")
        threshold_tddi = st.number_input("TDD I threshold (%):", value=5.0, step=0.5, format="%.1f")
        threshold_u_neg = st.number_input("u- threshold (%):", value=3.0, step=0.5, format="%.1f")
        threshold_vh = st.number_input("Voltage harmonic threshold (%):", value=3.0, step=0.1, format="%.1f")
        threshold_ch = st.number_input("Current harmonic threshold (%):", value=4.0, step=0.1, format="%.1f")
    
    st.markdown("---")
    st.markdown("#### Rated Power")
    PDM_default = 40000000

    Pdm_max = st.number_input("Pdm (W)", value=PDM_default, step=1000000, format="%d")
    
    st.markdown("---")
    run_button = st.button("▶️ START PROCESSING", use_container_width=True)
    
    if run_button:
        st.session_state.run_processing = True
        st.session_state.processing_complete = False
        st.session_state.processed_data = {}

# MAIN CONTENT AREA
st.markdown('<div class="main-title">⚡ POWER QUALITY ANALYSIS SYSTEM ⚡</div>', unsafe_allow_html=True)

# Processing section
if st.session_state.run_processing and not st.session_state.processing_complete:
    files_to_process = []
    if uploaded_files:
        for f in uploaded_files:
            files_to_process.append(("uploaded", f.name, f))
    elif folder_input:
        if os.path.isdir(folder_input):
            files = glob.glob(os.path.join(folder_input, "*.xlsx"))
            for f in files:
                files_to_process.append(("local", os.path.basename(f), f))
        else:
            st.error("❌ Invalid folder path")
            st.stop()
    else:
        st.error("❌ Please upload files or specify a folder path")
        st.stop()
        update_thresholds(threshold_pst, threshold_plt, threshold_thdu, threshold_tddi, threshold_u_neg)
    col1, col2 = st.columns(2)

    with col1:
        st.markdown(f"""
        <div class="stat-card">
            <div class="stat-number">{len(files_to_process)}</div>
            <div class="stat-label">Files to Process</div>
        </div>
        """, unsafe_allow_html=True)

    with col2:
        st.markdown(f"""
        <div class="stat-card">
            <div class="stat-number">{Pdm_max/1000000:.1f}</div>
            <div class="stat-label">Pdm (MW)</div>
        </div>
        """, unsafe_allow_html=True)


    st.markdown("---")
    progress_bar = st.progress(0)
    status_text = st.empty()
    tmp_out_dir = tempfile.mkdtemp(prefix="pq_out_")
    processed_count = 0

    for idx, (ftype, name, fobj) in enumerate(files_to_process, start=1):
        status_text.info(f"🔄 Processing ({idx}/{len(files_to_process)}): **{name}**")
        cleaning_log = ""
        try:
            if ftype == "uploaded":
                xls = pd.ExcelFile(fobj)
            else:
                xls = pd.ExcelFile(fobj)
            
            if "Measurements" in xls.sheet_names:
                sheet_to_read = "Measurements"
            else:
                sheet_to_read = None
                for sheet in xls.sheet_names:
                    try:
                        temp = pd.read_excel(xls, sheet_name=sheet, nrows=1)
                    except Exception:
                        continue
                    if any(col in temp.columns for col in cols_pst):
                        sheet_to_read = sheet
                        break
                if sheet_to_read is None:
                    status_text.warning(f"⚠️ {name}: No matching sheet found — skipping")
                    progress_bar.progress(int(idx/len(files_to_process)*100))
                    continue         
            data = pd.read_excel(xls, sheet_name=sheet_to_read)

            data_raw = data.copy()
            if enable_cleaning:
                time_col_candidates = ["Time [UTC]", "Time", "Datetime"]
                time_col_found = next((c for c in time_col_candidates if c in data.columns), None)
                if time_col_found:
                    status_text.info(f"🧹 Cleaning data for: **{name}**")
                    data, cleaning_log = clean_data(data, time_col_found)
                    if data.empty:
                        status_text.warning(f"⚠️ {name}: {cleaning_log} — skipping")
                        progress_bar.progress(int(idx/len(files_to_process)*100))
                        continue

            if "Time [UTC]" in data.columns:
                time_col = data["Time [UTC]"].copy()
                data = data.drop(columns=["Time [UTC]"])
            else:
                time_col = None
            data = data.replace("-", np.nan)
            data = data.apply(pd.to_numeric, errors="ignore")
            
            # BƯỚC 1: Lọc theo ngày đạt yêu cầu (có ít nhất 1 mẫu Ptot >= 50% Pdm)
            data_for_plt = None
            data_for_plt_6to18 = None
            
            if "Ptot+(Avg) [W]" in data.columns and time_col is not None:
                # Thêm cột Date và Time để phân loại
                data_temp = data.copy()
                data_temp["Date"] = pd.to_datetime(time_col.loc[data_temp.index], errors='coerce').dt.date
                data_temp["Time"] = pd.to_datetime(time_col.loc[data_temp.index], errors='coerce')
                data_temp = data_temp.dropna(subset=["Date"])
                
                if len(data_temp) > 0:
                    # Tính %Pdm cho mỗi mẫu
                    data_temp["%Pdm"] = (data_temp["Ptot+(Avg) [W]"] / Pdm_max * 100)
                    
                    # Tìm các ngày đạt yêu cầu (có ít nhất 1 mẫu >= 50% Pdm)
                    valid_dates = []
                    for date in data_temp["Date"].unique():
                        date_data = data_temp[data_temp["Date"] == date]
                        if (date_data["%Pdm"] >= 50).any():
                            valid_dates.append(date)
                    
                    # Lọc chỉ giữ lại dữ liệu của các ngày đạt yêu cầu
                    if len(valid_dates) > 0:
                        data_valid_days = data_temp[data_temp["Date"].isin(valid_dates)].copy()
                        status_text.info(f"📅 Found {len(valid_dates)} valid days (with Ptot ≥ 50% Pdm) out of {len(data_temp['Date'].unique())} total days")
                        
                        # Tạo data cho Plt (toàn bộ ngày)
                        data_for_plt = data_valid_days.drop(columns=["Date", "%Pdm", "Time"]).copy()
                        
                        # Tạo data cho Plt 6:00-18:00
                        data_valid_days["Hour"] = data_valid_days["Time"].dt.hour
                        data_6to18 = data_valid_days[(data_valid_days["Hour"] >= 6) & (data_valid_days["Hour"] < 18)].copy()
                        data_for_plt_6to18 = data_6to18.drop(columns=["Date", "%Pdm", "Time", "Hour"]).copy()
                        
                        # BƯỚC 2: Lọc thêm Ptot > 0 cho các bảng còn lại
                        data = data_valid_days.copy()
                        invalid_mask = (data["Ptot+(Avg) [W]"].isna()) | (data["Ptot+(Avg) [W]"] <= 0)
                        if invalid_mask.any():
                            data = data[~invalid_mask]
                        data = data.drop(columns=["Date", "%Pdm", "Time"], errors='ignore')
                        if "Hour" in data.columns:
                            data = data.drop(columns=["Hour"])
                    else:
                        status_text.warning(f"⚠️ {name}: No valid days found (no day with Ptot ≥ 50% Pdm) — skipping")
                        progress_bar.progress(int(idx/len(files_to_process)*100))
                        continue
            
            data = data.fillna(0)
            if time_col is not None:
                time_filtered = time_col.loc[data.index]
                data.insert(0, "Time [UTC]", time_filtered)

            stat_summary_pst = pd.DataFrame()
            stat_summary_plt = pd.DataFrame()
            stat_summary_plt_6to18 = pd.DataFrame()
            table_plt_24h = pd.DataFrame()
            table_plt_6to18 = pd.DataFrame()
            stat_summary_uneg = pd.DataFrame()
            stat_summary_vh = pd.DataFrame()
            stat_summary_ch = pd.DataFrame()
            daily_pdm_table = pd.DataFrame()
            daily_summary_text = ""
            pdm_detailed_info = {}

            
# TÍNH TOÁN PLT (ưu tiên, chỉ cần Ptot >= 50% Pdm)
            plt_valid_days_count = 0
            plt_total_samples = 0
            plt_divided_by_12 = 0
                        
            if data_for_plt is not None and all(c in data_for_plt.columns for c in cols_plt):
                total_samples_from_valid_days = len(data_for_plt)
                plt_total_samples = total_samples_from_valid_days
                plt_divided_by_12 = total_samples_from_valid_days / 12
                
                # Đếm số ngày đủ điều kiện cho Plt
                if "Date" in data_temp.columns:
                    plt_valid_days_count = len(valid_dates)
                
                total_plt_values = total_samples_from_valid_days / 12
                
                passed = {}
                failed = {}
                perc = {}
                status_per_col = {}
                
                for col in cols_plt:
                    threshold = thresholds[col]
                    pass_count = ((data_for_plt[col] <= threshold) & data_for_plt[col].notna()).sum() / 12
                    fail_count = total_plt_values - pass_count
                    percentage = (pass_count / total_plt_values * 100) if total_plt_values > 0 else 0
                    
                    passed[col] = pass_count
                    failed[col] = fail_count
                    perc[col] = round(percentage, 2)
                    status_per_col[col] = "PASS" if percentage >= 95 else "FAIL"
                
                group_status = "PASS" if all(s == "PASS" for s in status_per_col.values()) else "FAIL"
                
                table_plt_24h = pd.DataFrame([
                    {col: total_plt_values for col in cols_plt},
                    passed,
                    failed,
                    perc,
                    status_per_col
                ], index=[
                    "No. of samples",
                    "No. of pass samples",
                    "No. of fail samples",
                    "Percentage of pass samples [%]",
                    "Status"
                ])
                table_plt_24h.loc["Group Status"] = [group_status] * len(cols_plt)
                
                plt_sample_count = total_samples_from_valid_days / 12
                
                
                # Plt Statistics (Overall) cho toàn bộ ngày
                plt_stats = []
                for col in cols_plt:
                    mx = data_for_plt[col].max()
                    mn = data_for_plt[col].min()
                    avg = data_for_plt[col].mean()
                    plt_stats.extend([mx, avg, mn])
                
                plt_columns = pd.MultiIndex.from_tuples([
                    ("Phase A", "Max"), ("Phase A", "AVG"), ("Phase A", "Min"),
                    ("Phase B", "Max"), ("Phase B", "AVG"), ("Phase B", "Min"),
                    ("Phase C", "Max"), ("Phase C", "AVG"), ("Phase C", "Min")
                ])
                stat_summary_plt = pd.DataFrame([plt_stats], columns=plt_columns)
                stat_summary_plt.index = ["Overall (24h)"]
                stat_summary_plt.index.name = "Statistic"
                stat_summary_plt = stat_summary_plt.round(3)
            
            if data_for_plt_6to18 is not None and len(data_for_plt_6to18) > 0 and all(c in data_for_plt_6to18.columns for c in cols_plt):
                # TÍNH TOÁN BẢNG PLT 6-18H - Dựa trên tổng mẫu chia 12
                total_samples_6to18 = len(data_for_plt_6to18)
                total_plt_values_6to18 = total_samples_6to18 / 12
                
                passed_6to18 = {}
                failed_6to18 = {}
                perc_6to18 = {}
                status_per_col_6to18 = {}
                
                for col in cols_plt:
                    threshold = thresholds[col]
                    pass_count = ((data_for_plt_6to18[col] <= threshold) & data_for_plt_6to18[col].notna()).sum() / 12
                    fail_count = total_plt_values_6to18 - pass_count
                    percentage = (pass_count / total_plt_values_6to18 * 100) if total_plt_values_6to18 > 0 else 0
                    
                    passed_6to18[col] = pass_count
                    failed_6to18[col] = fail_count
                    perc_6to18[col] = round(percentage, 2)
                    status_per_col_6to18[col] = "PASS" if percentage >= 95 else "FAIL"
                
                group_status_6to18 = "PASS" if all(s == "PASS" for s in status_per_col_6to18.values()) else "FAIL"
                
                table_plt_6to18 = pd.DataFrame([
                    {col: total_plt_values_6to18 for col in cols_plt},
                    passed_6to18,
                    failed_6to18,
                    perc_6to18,
                    status_per_col_6to18
                ], index=[
                    "No. of samples",
                    "No. of pass samples",
                    "No. of fail samples",
                    "Percentage of pass samples [%]",
                    "Status"
                ])
                table_plt_6to18.loc["Group Status"] = [group_status_6to18] * len(cols_plt)

                
                # Plt Statistics (Overall) cho 6:00-18:00
                plt_stats_6to18 = []
                for col in cols_plt:
                    mx = data_for_plt_6to18[col].max()
                    mn = data_for_plt_6to18[col].min()
                    avg = data_for_plt_6to18[col].mean()
                    plt_stats_6to18.extend([mx, avg, mn])
                
                plt_columns = pd.MultiIndex.from_tuples([
                    ("Phase A", "Max"), ("Phase A", "AVG"), ("Phase A", "Min"),
                    ("Phase B", "Max"), ("Phase B", "AVG"), ("Phase B", "Min"),
                    ("Phase C", "Max"), ("Phase C", "AVG"), ("Phase C", "Min")
                ])
                stat_summary_plt_6to18 = pd.DataFrame([plt_stats_6to18], columns=plt_columns)
                stat_summary_plt_6to18.index = ["Overall (6:00-18:00)"]
                stat_summary_plt_6to18.index.name = "Statistic"
                stat_summary_plt_6to18 = stat_summary_plt_6to18.round(3)

            if "Ptot+(Avg) [W]" in data.columns:
                Pdm_max_val = Pdm_max
                if pd.isna(Pdm_max_val) or Pdm_max_val <= 0:
                    pass

                else:
                    daily_pdm_table, daily_summary_text, pdm_detailed_info = make_daily_pdm_distribution(data, Pdm_max_val, data_raw)                  
                    bins = []
                    labels = []
                    for i in range(1, 11):
                        lower = (i - 1) * 0.1 * Pdm_max_val
                        upper = i * 0.1 * Pdm_max_val
                        bins.append((lower, upper))
                        labels.append(f"{i*10}%")
                    
                    stat_summary_vh = make_voltage_harmonics_by_pdm(data, bins, labels)
                    stat_summary_ch = make_current_harmonics_by_pdm(data, bins, labels)
                    
                    rows_pst = []
                    rows_uneg = []
                    required_stat_cols = {"Pst": cols_pst, "Uneg": cols_uneg}
                    
                    for (lower, upper), label in zip(bins, labels):
                        if lower <= 0:
                            mask = (data["Ptot+(Avg) [W]"] > 0) & (data["Ptot+(Avg) [W]"] <= upper)
                        else:
                            mask = (data["Ptot+(Avg) [W]"] > lower) & (data["Ptot+(Avg) [W]"] <= upper)
                        df_bin = data.loc[mask]
                        count = int(mask.sum())
                        
                        pst_vals = []
                        for col in required_stat_cols["Pst"]:
                            if col in data.columns and count > 0:
                                mx = df_bin[col].max()
                                mn = df_bin[col].min()
                                avg = df_bin[col].mean()
                            else:
                                mx = np.nan
                                mn = np.nan
                                avg = np.nan
                            pst_vals.append((mx, avg, mn))
                        rows_pst.append([label] + [v for triple in pst_vals for v in triple])
                        
                        ucol = required_stat_cols["Uneg"][0]
                        if ucol in data.columns and count > 0:
                            mx = df_bin[ucol].max()
                            mn = df_bin[ucol].min()
                            avg = df_bin[ucol].mean()
                        else:
                            mx = np.nan
                            mn = np.nan
                            avg = np.nan
                        rows_uneg.append([label, mx, avg, mn])
                    
                    pst_columns = pd.MultiIndex.from_tuples([
                        ("Phase A", "Max"), ("Phase A", "AVG"), ("Phase A", "Min"),
                        ("Phase B", "Max"), ("Phase B", "AVG"), ("Phase B", "Min"),
                        ("Phase C", "Max"), ("Phase C", "AVG"), ("Phase C", "Min")
                    ])
                    pst_df = pd.DataFrame([r[1:] for r in rows_pst], index=[r[0] for r in rows_pst], columns=pst_columns)
                    pst_df.index.name = "%Pdm"
                    
                    uneg_df = pd.DataFrame([r[1:] for r in rows_uneg], index=[r[0] for r in rows_uneg], columns=["Max", "AVG", "Min"])
                    uneg_df.index.name = "%Pdm"
                    
                    stat_summary_pst = pst_df.round(3)
                    stat_summary_uneg = uneg_df.round(3)

            # TÍNH TOÁN CÁC BẢNG CÒN LẠI (sau khi lọc Ptot > 0)
            table_pst = make_table(data, cols_pst) if all(c in data.columns for c in cols_pst) else pd.DataFrame()
            table_thdu = make_table(data, cols_thdu) if all(c in data.columns for c in cols_thdu) else pd.DataFrame()
            table_tddi = make_table(data, cols_tddi) if all(c in data.columns for c in cols_tddi) else pd.DataFrame()
            table_uneg = make_table(data, cols_uneg) if all(c in data.columns for c in cols_uneg) else pd.DataFrame()
            table_vh_order = make_voltage_harmonics_detail(data, threshold_vh)
            table_ch_order = make_current_harmonics_detail(data, threshold_ch)
            table_thd_max = make_thd_max_table(data)
            table_tdd_max = make_tdd_max_table(data)
            table_vh_max = make_voltage_harmonics_max_table(data)
            table_ch_max = make_current_harmonics_max_table(data)
            if use_local_save and out_folder_input:
                save_folder = out_folder_input
                os.makedirs(save_folder, exist_ok=True)
            elif use_local_save and folder_input:
                save_folder = os.path.join(folder_input, "Output")
                os.makedirs(save_folder, exist_ok=True)
            else:
                save_folder = tmp_out_dir
            
            base_name = os.path.splitext(name)[0]
            out_excel_path = os.path.join(save_folder, f"Output_{base_name}.xlsx")
            out_png_path = os.path.join(save_folder, f"Output_{base_name}.png")
            out_pdf_path = os.path.join(save_folder, f"Output_{base_name}.pdf")

            with pd.ExcelWriter(out_excel_path, engine="openpyxl") as writer:
                if not table_pst.empty:
                    save_table_to_excel(writer, table_pst, "Pst")
                if not table_thdu.empty:
                    save_table_to_excel(writer, table_thdu, "THDu")
                if not table_tddi.empty:
                    save_table_to_excel(writer, table_tddi, "TDDi")
                
                # Plt tables (24h and 6-18)
                if not table_plt_24h.empty:
                    save_table_to_excel(writer, table_plt_24h, "Plt_24h")
                if not table_plt_6to18.empty:
                    save_table_to_excel(writer, table_plt_6to18, "Plt_6to18")
                
                if not table_uneg.empty:
                    save_table_to_excel(writer, table_uneg, "u0Avg")
                table_vh_order.to_excel(writer, sheet_name="Voltage Harmonics by Order")
                table_ch_order.to_excel(writer, sheet_name="Current Harmonics by Order")
                if not table_thd_max.empty:
                    table_thd_max.to_excel(writer, sheet_name="THD_Max_Values", index=False)
                if not table_tdd_max.empty:
                    table_tdd_max.to_excel(writer, sheet_name="TDD_Max_Values", index=False)
                if not daily_pdm_table.empty:
                    daily_pdm_table.to_excel(writer, sheet_name="Daily_Pdm_Distribution", index=False)
                if not table_vh_max.empty:
                    table_vh_max.to_excel(writer, sheet_name="Voltage_Harmonics_Max", index=False)
                if not table_ch_max.empty:
                    table_ch_max.to_excel(writer, sheet_name="Current_Harmonics_Max", index=False)
                
                try:
                    if not stat_summary_pst.empty:
                        stat_summary_pst.to_excel(writer, sheet_name="Power_Stats_Pst", startrow=1)
                    if not stat_summary_plt.empty:
                        stat_summary_plt.to_excel(writer, sheet_name="Plt_Statistics_24h", startrow=1)
                    if not stat_summary_plt_6to18.empty:
                        stat_summary_plt_6to18.to_excel(writer, sheet_name="Plt_Statistics_6to18", startrow=1)
                    if not stat_summary_uneg.empty:
                        stat_summary_uneg.to_excel(writer, sheet_name="Power_Stats_Uneg", startrow=1)
                    if not stat_summary_vh.empty:
                        stat_summary_vh.to_excel(writer, sheet_name="Voltage_Harmonics_by_%Pdm", startrow=1)
                    if not stat_summary_ch.empty:
                        stat_summary_ch.to_excel(writer, sheet_name="Current_Harmonics_by_%Pdm", startrow=1)
                except Exception:
                    pass

            png_saved = plot_and_save(data, base_name, save_folder)
            
            tables_for_pdf = {
                "Pst": table_pst,
                "THDu": table_thdu,
                "TDDi": table_tddi,
                "THD Maximum Values": table_thd_max, 
                "TDD Maximum Values": table_tdd_max,
                "Plt (24h)": table_plt_24h,
                "Plt (6:00-18:00)": table_plt_6to18,
                "u0Avg": table_uneg,
                "Voltage Harmonics by Order": table_vh_order,
                "Current Harmonics by Order": table_ch_order,
                "Voltage Harmonics Max Values": table_vh_max,
                "Current Harmonics Max Values": table_ch_max
            }
            if not daily_pdm_table.empty:
                tables_for_pdf["Daily Pdm Distribution"] = daily_pdm_table
            if not stat_summary_pst.empty:
                tables_for_pdf["Pst by %Pdm"] = stat_summary_pst
            if not stat_summary_plt.empty:
                tables_for_pdf["Plt Statistics (24h)"] = stat_summary_plt
            if not stat_summary_plt_6to18.empty:
                tables_for_pdf["Plt Statistics (6:00-18:00)"] = stat_summary_plt_6to18
            if not stat_summary_uneg.empty:
                tables_for_pdf["Uneg by %Pdm"] = stat_summary_uneg
            if not stat_summary_vh.empty:
                tables_for_pdf["Voltage Harmonics (2-40) by %Pdm"] = stat_summary_vh
            if not stat_summary_ch.empty:
                tables_for_pdf["Current Harmonics (2-40) by %Pdm"] = stat_summary_ch

            save_pdf(tables_for_pdf, out_pdf_path, base_name)
            word_report_path = None
            word_report_log = ""
            if pdm_detailed_info:
                valid_samples_count = pdm_detailed_info.get('valid_samples', len(data))
                total_samples_count = pdm_detailed_info.get('total_samples', len(data))
            else:
                if data_raw is not None and "Ptot+(Avg) [W]" in data_raw.columns:
                    total_samples_count = len(data_raw)
                    ptot_values = pd.to_numeric(data_raw["Ptot+(Avg) [W]"], errors='coerce')
                    valid_samples_count = len(data_raw[ptot_values > 0])
                else:
                    total_samples_count = len(data)
                    valid_samples_count = len(data)
            
            st.session_state.processed_data[base_name] = {
                'data': data.copy(),
                'data_raw': data_raw,
                'voltage_level': voltage_level,
                'thresholds': {
                    'pst': threshold_pst,
                    'plt': threshold_plt,
                    'thdu': threshold_thdu,
                    'tddi': threshold_tddi,
                    'u_neg': threshold_u_neg,
                    'vh': threshold_vh,
                    'ch': threshold_ch,
                    'pdm_mw': Pdm_max / 1_000_000,  # Convert W to MW
                    'valid_samples': valid_samples_count, 
                    'total_samples': total_samples_count  
                },
                'tables': {

                    'table_pst': table_pst,
                    'table_thdu': table_thdu,
                    'table_tddi': table_tddi,
                    'table_plt_24h': table_plt_24h,
                    'table_plt_6to18': table_plt_6to18,
                    'table_uneg': table_uneg,
                    'table_vh_order': table_vh_order,
                    'table_ch_order': table_ch_order,
                    'table_thd_max': table_thd_max, 
                    'table_tdd_max': table_tdd_max,
                    'stat_summary_pst': stat_summary_pst,
                    'stat_summary_plt': stat_summary_plt,
                    'stat_summary_plt_6to18': stat_summary_plt_6to18,
                    'stat_summary_uneg': stat_summary_uneg,
                    'stat_summary_vh': stat_summary_vh,
                    'stat_summary_ch': stat_summary_ch,
                    'table_vh_max': table_vh_max,
                    'table_ch_max': table_ch_max,
                    'daily_pdm_table': daily_pdm_table,
                    'daily_summary_text': daily_summary_text,
                    'pdm_detailed_info': pdm_detailed_info 

                },
                'excel_path': out_excel_path,
                'png_path': out_png_path,
                'pdf_path': out_pdf_path,
                'save_folder': save_folder,
                'plt_info': {
                    'valid_days_count': plt_valid_days_count,
                    'total_samples': plt_total_samples,
                    'divided_by_12': plt_divided_by_12
                },
                'cleaning_log': cleaning_log,
                'word_report_path': word_report_path,
                'word_report_log': word_report_log
            }
            processed_count += 1
        except Exception as e:
            st.error(f"❌ Error processing {name}: {e}")
        progress_bar.progress(int(idx/len(files_to_process)*100))

    st.session_state.processing_complete = True
    st.session_state.run_processing = False
    status_text.success(f"✅ COMPLETED: Successfully processed {processed_count}/{len(files_to_process)} file(s)")

# Display results section
if st.session_state.processing_complete and st.session_state.processed_data:
    st.markdown("---")
    
    # File selector with better styling
    file_names = list(st.session_state.processed_data.keys())
    
    col1, col2 = st.columns([3, 1])
    with col1:
        st.markdown("### 📊 Analysis Results")
    with col2:
        selected_file = st.selectbox("Select file:", file_names, label_visibility="collapsed")
    
    if selected_file:
        result = st.session_state.processed_data[selected_file]
        data = result['data']
        tables = result['tables']
        # Display voltage level and thresholds info
        st.markdown("---")
        st.markdown("### ⚙️ Analysis Configuration")
        col1, col2 = st.columns(2)
        with col1:
            st.markdown(f"""
            <div class="stat-card">
                <div class="stat-number">{result.get('voltage_level', 'N/A')}</div>
                <div class="stat-label">Voltage Level</div>
            </div>
            """, unsafe_allow_html=True)
        
        with col2:
            thresholds_info = result.get('thresholds', {})
            st.markdown(f"""
            <div class="stat-card" style="text-align: left; padding: 1rem 1.5rem;">
                <div style="font-weight: 600; margin-bottom: 0.5rem;">Applied Thresholds:</div>
                <div style="font-size: 0.9rem; line-height: 1.8;">
                    • Pst: {thresholds_info.get('pst', 'N/A')}<br/>
                    • Plt: {thresholds_info.get('plt', 'N/A')}<br/>
                    • THD U: {thresholds_info.get('thdu', 'N/A')}%<br/>
                    • TDD I: {thresholds_info.get('tddi', 'N/A')}%<br/>
                    • u_neg: {thresholds_info.get('u_neg', 'N/A')}%<br/>
                    • Voltage Harmonic: {thresholds_info.get('vh', 'N/A')}%<br/>
                    • Current Harmonic: {thresholds_info.get('ch', 'N/A')}%
                </div>
            </div>
            """, unsafe_allow_html=True)
        
        # Download section with cards
        # st.markdown('<div class="download-section">', unsafe_allow_html=True)
        st.markdown(f"#### 📥 Download Results for: **{selected_file}**")
        
        import base64
        def get_download_link(file_path, filename, label, icon):
            with open(file_path, "rb") as f:
                file_bytes = f.read()
            b64 = base64.b64encode(file_bytes).decode()
            if filename.endswith('.xlsx'):
                mime = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            elif filename.endswith('.png'):
                mime = "image/png"
            elif filename.endswith('.pdf'):
                mime = "application/pdf"
            else:
                mime = "application/octet-stream"
            href = f'<a href="data:{mime};base64,{b64}" download="{filename}" target="_blank" style="text-decoration:none;"><button style="background: linear-gradient(90deg, #3498db 0%, #2980b9 100%); color: white; border: none; border-radius: 8px; padding: 0.75rem 1.5rem; font-weight: 600; cursor: pointer; width: 100%; box-shadow: 0px 4px 10px rgba(52, 152, 219, 0.3);">{icon} {label}</button></a>'
            return href

        c1, c2, c3 = st.columns(3)
        c1.markdown(get_download_link(result['excel_path'], f"Output_{selected_file}.xlsx", "Excel Report", "📘"), unsafe_allow_html=True)
        c2.markdown(get_download_link(result['png_path'], f"Output_{selected_file}.png", "Charts (PNG)", "🖼️"), unsafe_allow_html=True)
        c3.markdown(get_download_link(result['pdf_path'], f"Output_{selected_file}.pdf", "PDF Report", "📄"), unsafe_allow_html=True)
        
        st.markdown('</div>', unsafe_allow_html=True)
# ==================== WORD REPORT GENERATOR SECTION ====================
        st.markdown("---")
        st.markdown("#### 📝 Word Report Generator")

        st.markdown("##### 📤 Upload Word Template")
        
        word_template_upload = st.file_uploader(
            "Select your Word template file (.docx)",
            type=["docx"],
            key=f"word_upload_{selected_file}",
            help="This template will be populated with data from the Excel report"
        )
        
        if word_template_upload:
            st.success(f"✅ Template uploaded: **{word_template_upload.name}**")
            
            if st.button("📝 Generate Word Report", 
                        type="primary", 
                        use_container_width=True,
                        key=f"generate_word_btn_{selected_file}"):
                
                with st.spinner("📝 Generating Word report..."):
                    try:
                        # Save uploaded template to temp file
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_template:
                            tmp_template.write(word_template_upload.getvalue())
                            temp_template_path = tmp_template.name
                        
                        # Output path
                        word_output_path = os.path.join(
                            result['save_folder'], 
                            f"Word_Report_{selected_file}.docx"
                        )
                        
                        # Generate Word report
                        success, log = process_word_report(
                            result['excel_path'], 
                            temp_template_path, 
                            word_output_path,
                            result.get('thresholds', {}) 
                        )
                        
                        # Clean up temp file
                        try:
                            os.unlink(temp_template_path)
                        except:
                            pass
                        
                        # Display results
                        if success:
                            st.success("✅ Word report generated successfully!")
                            
                            # Show log in expander
                            with st.expander("📋 View Processing Log", expanded=False):
                                st.text(log)
                            
                            # Show file info
                            file_size = os.path.getsize(word_output_path) / 1024  # KB
                            st.info(f"📄 File size: {file_size:.2f} KB")
                            
                            # Download button
                            with open(word_output_path, "rb") as f:
                                st.download_button(
                                    label="📥 Download Word Report",
                                    data=f,
                                    file_name=f"Word_Report_{selected_file}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    use_container_width=True,
                                    type="primary",
                                    key=f"download_word_{selected_file}"
                                )
                        else:
                            st.error("❌ Failed to generate Word report")
                            with st.expander("📋 View Error Log", expanded=True):
                                st.text(log)
                    
                    except Exception as e:
                        st.error(f"❌ Error: {str(e)}")
                        import traceback
                        with st.expander("📋 View Detailed Error", expanded=True):
                            st.code(traceback.format_exc())
        else:
            st.info("👆 Please upload a Word template (.docx) to generate the report")

        st.markdown('</div>', unsafe_allow_html=True)
        
# Summary statistics
        if tables['daily_summary_text']:
            st.markdown("---")
            if result.get('cleaning_log'):
                with st.expander("🧹 Data Cleaning Log", expanded=False):
                    st.text(result['cleaning_log'])
            st.info(tables['daily_summary_text'])

            # ✅ HIỂN THỊ THÔNG TIN PLT
            if 'plt_info' in result and result['plt_info']['total_samples'] > 0:
                plt_info = result['plt_info']
                remainder = plt_info['total_samples'] % 12
                
                st.markdown("### 📊 Plt Calculation Summary")
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    st.markdown(f"""
                    <div class="stat-card">
                        <div class="stat-number">{plt_info['valid_days_count']}</div>
                        <div class="stat-label">Valid Days for Plt<br/>(Ptot≥50%Pdm)</div>
                    </div>
                    """, unsafe_allow_html=True)
                
                with col2:
                    st.markdown(f"""
                    <div class="stat-card">
                        <div class="stat-number">{plt_info['total_samples']:,}</div>
                        <div class="stat-label">Total Samples (Ptot≥50%Pdm)</div>
                    </div>
                    """, unsafe_allow_html=True)
                
                with col3:
                    if remainder == 0:
                        color = "#27ae60"
                        status = "✅"
                    else:
                        color = "#e74c3c"
                        status = f"⚠️: {remainder}"
                    
                    st.markdown(f"""
                    <div class="stat-card">
                        <div class="stat-number" style="color: {color};">{plt_info['divided_by_12']:.2f}</div>
                        <div class="stat-label">Plt Samples <br/>{status}</div>
                    </div>
                    """, unsafe_allow_html=True)

        if 'pdm_detailed_info' in tables and tables['pdm_detailed_info']:
            st.markdown("---")            
            info = tables['pdm_detailed_info']
            
            # Statistics cards
            col1, col2, col3 = st.columns(3)
            with col1:
                st.markdown(f"""
                <div class="stat-card">
                    <div class="stat-number">{info['valid_days_count']}</div>
                    <div class="stat-label">Valid Days</div>
                </div>
                """, unsafe_allow_html=True)
            
            with col2:
                st.markdown(f"""
                <div class="stat-card">
                    <div class="stat-number">{info['total_valid_days_samples']:,}</div>
                    <div class="stat-label">Total Samples (All Valid Days)</div>
                </div>
                """, unsafe_allow_html=True)
            
            with col3:
                avg_samples = info['total_valid_days_samples'] / info['valid_days_count'] if info['valid_days_count'] > 0 else 0
                st.markdown(f"""
                <div class="stat-card">
                    <div class="stat-number">{avg_samples:.0f}</div>
                    <div class="stat-label">Avg Samples/Day</div>
                </div>
                """, unsafe_allow_html=True)
            
            # Detailed table with expander
            with st.expander("📊 View Detailed Breakdown by Date", expanded=False):
                valid_dates_df = pd.DataFrame([
                    {
                        'No.': idx,
                        'Date': date.strftime("%d/%m/%Y"),
                        'Samples': samples,
                        # 'Status': '✅ Valid (≥50% Pdm)'
                    }
                    for idx, (date, samples) in enumerate(sorted(info['valid_dates_samples'].items()), 1)
                ])
                
                # Add total row
                total_row = pd.DataFrame([{
                    'No.': '',
                    'Date': 'TOTAL',
                    'Samples': info['total_valid_days_samples'],
                    'Status': f"{info['valid_days_count']} days"
                }])
                valid_dates_df = pd.concat([valid_dates_df, total_row], ignore_index=True)
                
                st.dataframe(valid_dates_df, use_container_width=True, hide_index=True)

        st.markdown("---")
        tab1, tab2, tab3, tab4, tab5 = st.tabs(["📋 Summary Tables", "📊 Harmonics Details", "⚡ Power Statistics", "📈 Raw Data", "🎨 Custom Plot"])
        
        with tab1:
            col1, col2 = st.columns(2)
            with col1:
                if not tables['table_pst'].empty:
                    st.markdown("##### Pst Summary")
                    st.dataframe(tables['table_pst'], use_container_width=True)
                if not tables['table_plt_24h'].empty:
                    st.markdown("##### Plt Summary (24h) - Divided by 12")
                    st.dataframe(tables['table_plt_24h'], use_container_width=True)
                if not tables['table_plt_6to18'].empty:
                    st.markdown("##### Plt Summary (6:00-18:00) - Divided by 12")
                    st.dataframe(tables['table_plt_6to18'], use_container_width=True)
            with col2:
                if not tables['table_thdu'].empty:
                    st.markdown("##### THD U Summary")
                    st.dataframe(tables['table_thdu'], use_container_width=True)
                if not tables['table_tddi'].empty:
                    st.markdown("##### TDD I Summary")
                    st.dataframe(tables['table_tddi'], use_container_width=True)
            
            if not tables['table_uneg'].empty:
                st.markdown("##### Uneg (u0Avg) Summary")
                st.dataframe(tables['table_uneg'], use_container_width=True)
        
            col1, col2 = st.columns(2)
            with col1:
                if not tables.get('table_thd_max', pd.DataFrame()).empty:
                    st.markdown("##### THD Maximum Values")
                    st.dataframe(tables['table_thd_max'], use_container_width=True, hide_index=True)
            with col2:
                if not tables.get('table_tdd_max', pd.DataFrame()).empty:
                    st.markdown("##### TDD Maximum Values")
                    st.dataframe(tables['table_tdd_max'], use_container_width=True, hide_index=True)
            col1, col2 = st.columns(2)
            with col1:
                if not tables.get('table_vh_max', pd.DataFrame()).empty:
                    st.markdown("##### Voltage Harmonics Max Values (Order 2-40)")
                    st.dataframe(tables['table_vh_max'], use_container_width=True, hide_index=True)
            with col2:
                if not tables.get('table_ch_max', pd.DataFrame()).empty:
                    st.markdown("##### Current Harmonics Max Values (Order 2-40)")
                    st.dataframe(tables['table_ch_max'], use_container_width=True, hide_index=True)

        with tab2:
            col1, col2 = st.columns(2)
            with col1:
                if not tables['table_vh_order'].empty:
                    st.markdown("##### Voltage Harmonics (Order 2-40)")
                    st.dataframe(tables['table_vh_order'], use_container_width=True, height=500)
            with col2:
                if not tables['table_ch_order'].empty:
                    st.markdown("##### Current Harmonics (Order 2-40)")
                    st.dataframe(tables['table_ch_order'], use_container_width=True, height=500)
        
        with tab3:
            if not tables['daily_pdm_table'].empty:
                st.markdown("##### Daily %Pdm Distribution")
                st.dataframe(tables['daily_pdm_table'], use_container_width=True)
            
            col1, col2 = st.columns(2)
            with col1:
                if not tables['stat_summary_pst'].empty:
                    st.markdown("##### Pst by %Pdm")
                    st.dataframe(tables['stat_summary_pst'], use_container_width=True)
            with col2:
                if not tables['stat_summary_uneg'].empty:
                    st.markdown("##### Uneg by %Pdm")
                    st.dataframe(tables['stat_summary_uneg'], use_container_width=True)
            
            if not tables['stat_summary_plt'].empty:
                st.markdown("##### Plt Statistics (24h - Overall)")
                st.dataframe(tables['stat_summary_plt'], use_container_width=True)
            
            if not tables['stat_summary_plt_6to18'].empty:
                st.markdown("##### Plt Statistics (6:00-18:00 - Overall)")
                st.dataframe(tables['stat_summary_plt_6to18'], use_container_width=True)
            
            if not tables['stat_summary_vh'].empty:
                st.markdown("##### Voltage Harmonics by %Pdm")
                st.dataframe(tables['stat_summary_vh'], use_container_width=True, height=400)
            
            if not tables['stat_summary_ch'].empty:
                st.markdown("##### Current Harmonics by %Pdm")
                st.dataframe(tables['stat_summary_ch'], use_container_width=True, height=400)
        
    with tab4:
        # Thêm radio button để chọn
        view_mode = st.radio("View mode:", ["Processed Data", "Raw Data (Original)"], horizontal=True)
        
        if view_mode == "Raw Data (Original)":
            display_data_source = result.get('data_raw', data)
            st.info("📋 Showing original data before any processing")
        else:
            display_data_source = data
            st.info("📋 Showing processed data (filtered by valid days and Ptot > 0)")
        
        st.markdown(f"**Total rows:** {len(display_data_source)} | **Total columns:** {len(display_data_source.columns)}")
        
        col1, col2, col3 = st.columns(3)
        with col1:
            show_rows = st.number_input("Rows to display:", min_value=10, max_value=len(display_data_source), value=min(100, len(display_data_source)), step=10)
        with col2:
            start_row = st.number_input("Start from row:", min_value=0, max_value=max(0, len(display_data_source)-1), value=0, step=10)
        with col3:
            search_col = st.text_input("Filter columns:", value="")
        
        if search_col.strip():
            search_terms = [term.strip() for term in search_col.split(",")]
            filtered_cols = [col for col in display_data_source.columns if any(term.lower() in col.lower() for term in search_terms)]
            if filtered_cols:
                display_data = display_data_source[filtered_cols].iloc[start_row:start_row + show_rows]
                st.success(f"✓ Found {len(filtered_cols)} matching columns")
            else:
                display_data = display_data_source.iloc[start_row:start_row + show_rows]
                st.warning("No matching columns found")
        else:
            display_data = display_data_source.iloc[start_row:start_row + show_rows]
        
        st.dataframe(display_data, use_container_width=True, height=400)
        
        if st.checkbox("Show statistics"):
            st.dataframe(display_data_source.describe(), use_container_width=True)
        
        with tab5:
            available_columns = list(data.columns)
            
            if len(available_columns) >= 2:
                col_x, col_y = st.columns(2)
                with col_x:
                    x_col = st.selectbox("X-axis:", available_columns, index=0)
                with col_y:
                    y_cols = st.multiselect(
                        "Y-axis (multiple):",
                        [c for c in available_columns if c != x_col],
                        default=[available_columns[1]] if len(available_columns) > 1 else []
                    )

                plot_type = st.radio("Chart type:", ["Line", "Scatter"], horizontal=True)
                
                if st.button("🎨 Generate Plot", use_container_width=True) and y_cols:
                    try:
                        import plotly.express as px
                        plot_df = data[[x_col] + y_cols].dropna()
                        max_points = 8000
                        if len(plot_df) > max_points:
                            plot_df = plot_df.sample(max_points).sort_values(by=x_col)
                            st.info(f"📊 Sampled to {max_points} points for performance")
                        
                        if plot_type == "Line":
                            fig = px.line(plot_df, x=x_col, y=y_cols)
                        else:
                            fig = px.scatter(plot_df, x=x_col, y=y_cols, opacity=0.7)
                        
                        fig.update_layout(
                            template="plotly_white",
                            height=500,
                            legend=dict(orientation="h", y=-0.2),
                            margin=dict(l=50, r=30, t=50, b=50)
                        )
                        st.plotly_chart(fig, use_container_width=True)
                    except Exception as e:
                        st.error(f"Error: {e}")
                elif y_cols == []:
                    st.warning("⚠️ Select at least one Y-axis column")
            else:
                st.info("Not enough columns for plotting")

elif not st.session_state.processing_complete:
    # Welcome screen
    st.markdown("""
    <style>
    @keyframes typing {
    from { width: 0 }
    to { width: 100% }
    }
    @keyframes blink {
    50% { border-color: transparent }
    }
    .typewriter h2 {
    overflow: hidden;
    border-right: .15em solid #00BCD4;
    white-space: nowrap;
    margin: 0 auto;
    letter-spacing: .05em;
    animation:
        typing 3.5s steps(40, end),
        blink .75s step-end infinite;
    }
    </style>

    <div style="text-align: center; padding: 3rem;">
    <div class="typewriter">
        <h2>Welcome to Power Quality Analysis System</h2>
    </div>
    <p style="font-size: 1.2rem; color: #7f8c8d; margin-top: 1rem;">
        Upload your Excel files or specify a folder path in the sidebar to begin analysis
    </p>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("""
    <style>
    .stat-card {
        transition: all 0.3s ease;
        border-radius: 15px;
        padding: 20px;
        background: #ffffff;
        box-shadow: 0 4px 10px rgba(0,0,0,0.05);
        cursor: pointer;
    }
    .stat-card:hover {
        transform: scale(1.05);
        box-shadow: 0 8px 20px rgba(0,0,0,0.15);
    }
    .stat-card:active {
        transform: scale(1.1);
    }
    </style>
    """, unsafe_allow_html=True)

    col1, col2, col3 = st.columns(3)
    with col1:
        st.markdown("""
        <div class="stat-card" style="text-align: center;">
            <div style="font-size: 3rem;">📊</div>
            <div style="font-weight: 600; margin-top: 1rem;">Comprehensive Analysis</div>
            <div style="color: #7f8c8d; margin-top: 0.5rem;">Analyze Pst, Plt, THD, TDD and harmonics</div>
        </div>
        """, unsafe_allow_html=True)

    with col2:
        st.markdown("""
        <div class="stat-card" style="text-align: center;">
            <div style="font-size: 3rem;">⚡</div>
            <div style="font-weight: 600; margin-top: 1rem;">Power Statistics</div>
            <div style="color: #7f8c8d; margin-top: 0.5rem;">Detailed %Pdm distribution analysis</div>
        </div>
        """, unsafe_allow_html=True)

    with col3:
        st.markdown("""
        <div class="stat-card" style="text-align: center;">
            <div style="font-size: 3rem;">📥</div>
            <div style="font-weight: 600; margin-top: 1rem;">Export Results</div>
            <div style="color: #7f8c8d; margin-top: 0.5rem;">Download Excel, PDF, and PNG reports</div>
        </div>
        """, unsafe_allow_html=True)
        
        